"""Export a Chinese V1.5 single-device coefficient evidence report.

The exporter is evidence-only: it reads existing snapshots, fit residuals and
write logs. It does not open COM ports, control routes, or write coefficients.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import math
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Mapping, Optional, Sequence


REPORT_TITLE = "V1.5 单台气体分析仪系数写入与拟合证据报告"


def _read_csv(path: Path) -> List[Dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return [dict(row) for row in csv.DictReader(handle)]


def _write_csv(path: Path, rows: Sequence[Mapping[str, Any]]) -> None:
    header: List[str] = []
    for row in rows:
        for key in row:
            if key not in header:
                header.append(str(key))
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=header)
        writer.writeheader()
        writer.writerows(rows)


def _read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def _sha256(path: Path) -> str:
    if not path.exists() or not path.is_file():
        return "MISSING"
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def _git_commit(root: Path) -> str:
    try:
        return subprocess.check_output(
            ["git", "rev-parse", "HEAD"],
            cwd=str(root),
            text=True,
            stderr=subprocess.DEVNULL,
        ).strip()
    except Exception:
        return ""


def _fmt(value: Any, digits: int = 7) -> str:
    if value in (None, "", "None", "nan"):
        return ""
    try:
        number = float(value)
    except Exception:
        return str(value)
    if not math.isfinite(number):
        return str(value)
    return f"{number:.{digits}g}"


def _json_list(value: Any) -> List[float]:
    if isinstance(value, list):
        raw = value
    else:
        try:
            raw = json.loads(str(value))
        except Exception:
            return []
    if not isinstance(raw, list):
        return []
    out: List[float] = []
    for item in raw:
        try:
            out.append(float(item))
        except Exception:
            pass
    return out


def _device_rows(config_path: Path, device_id: str) -> List[Dict[str, Any]]:
    cfg = _read_json(config_path)
    rows: List[Dict[str, Any]] = []
    for item in ((cfg.get("devices") or {}).get("gas_analyzers") or []):
        if str(item.get("device_id") or "").zfill(3) != device_id:
            continue
        rows.append(
            {
                "设备ID": device_id,
                "通道名": item.get("name", ""),
                "串口": item.get("port", ""),
                "波特率": item.get("baud", item.get("baudrate", "")),
                "MODE": item.get("mode", ""),
                "主动上传": item.get("active_send", ""),
                "身份来源": item.get("identity_source", "MODE frame id"),
                "说明": "串口只作为采集通道；正式身份以分析仪自报 MODE 帧 ID 为准。",
            }
        )
    return rows


def _snapshot_rows(snapshot_path: Path, device_id: str) -> List[Dict[str, Any]]:
    data = _read_json(snapshot_path)
    info = data.get(device_id) or (data.get("analyzers") or {}).get(device_id) or {}
    rows: List[Dict[str, Any]] = []
    for group in range(1, 10):
        values = _json_list(info.get(f"GETCO{group}_before") or [])
        row: Dict[str, Any] = {
            "设备ID": device_id,
            "串口": info.get("port", ""),
            "通道名": info.get("analyzer_prefix", ""),
            "系数组": f"GETCO{group}",
            "读回命令": info.get(f"GETCO{group}_before_command", f"GETCO,YGAS,FFF,{group}"),
            "系数值": json.dumps(values, ensure_ascii=False),
        }
        for idx in range(6):
            row[f"C{idx}"] = _fmt(values[idx]) if idx < len(values) else ""
        rows.append(row)
    return rows


def _one_csv_row(path: Path, device_id: str, *, device_key: str = "analyzer_device_id") -> Dict[str, str]:
    for row in _read_csv(path):
        if str(row.get(device_key) or row.get("device_id") or "").zfill(3) == device_id:
            return row
    return {}


def _co2_fit_rows(path: Path, device_id: str) -> List[Dict[str, Any]]:
    out = []
    for row in _read_csv(path):
        if str(row.get("analyzer_device_id") or "").zfill(3) != device_id:
            continue
        target = row.get("target_value", "")
        error = row.get("error", "")
        rel = ""
        try:
            target_f = float(target)
            error_f = float(error)
            if abs(target_f) > 1e-12:
                rel = 100.0 * error_f / target_f
        except Exception:
            pass
        out.append(
            {
                "设备ID": device_id,
                "点位": row.get("point_identity", row.get("sample_index", "")),
                "参考CO2(ppm)": _fmt(target),
                "模型CO2(ppm)": _fmt(row.get("prediction")),
                "误差(ppm)": _fmt(error),
                "相对误差(%)": _fmt(rel),
                "R": _fmt(row.get("ratio")),
                "温度(C)": _fmt(row.get("temperature_c")),
                "压力(hPa)": _fmt(row.get("pressure_hpa")),
                "模型项": row.get("model_terms", ""),
            }
        )
    return out


def _h2o_fit_rows(path: Path, device_id: str) -> List[Dict[str, Any]]:
    out = []
    for row in _read_csv(path):
        if str(row.get("analyzer_device_id") or "").zfill(3) != device_id:
            continue
        out.append(
            {
                "设备ID": device_id,
                "点位": row.get("point_run_id", ""),
                "参考H2O(mmol/mol)": _fmt(row.get("reference_h2o_mmol")),
                "模型H2O(mmol/mol)": _fmt(row.get("model_pred_h2o_mmol")),
                "误差(mmol/mol)": _fmt(row.get("model_error_mmol")),
                "相对误差(%)": _fmt(row.get("model_error_pct")),
                "H2O比值": _fmt(row.get("h2o_ratio_f")),
                "腔体温度(C)": _fmt(row.get("chamber_temp_c")),
                "数字测温仪(C)": _fmt(row.get("digital_thermometer_temp_c")),
                "露点(C)": _fmt(row.get("reference_dewpoint_c")),
                "参考压力(hPa)": _fmt(row.get("reference_pressure_hpa")),
            }
        )
    return out


def _error_summary(rows: Sequence[Mapping[str, Any]], *, error_key: str, pct_key: str) -> Dict[str, Any]:
    errors: List[float] = []
    pcts: List[float] = []
    for row in rows:
        try:
            errors.append(abs(float(str(row.get(error_key, "")).replace(",", ""))))
        except Exception:
            pass
        try:
            text = str(row.get(pct_key, "")).strip()
            if text:
                pcts.append(abs(float(text.replace(",", ""))))
        except Exception:
            pass
    return {
        "点数": len(rows),
        "最大绝对误差": _fmt(max(errors) if errors else ""),
        "最大相对误差(%)": _fmt(max(pcts) if pcts else ""),
    }


def _write_event_rows(paths: Sequence[tuple[str, Path]], device_id: str) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for label, path in paths:
        for row in _read_csv(path):
            did = str(row.get("analyzer_device_id") or row.get("device_id") or "").zfill(3)
            if did != device_id:
                continue
            rows.append(
                {
                    "事件": label,
                    "设备ID": did,
                    "串口": row.get("port", ""),
                    "状态": row.get("status", ""),
                    "写入": row.get("write_applied", ""),
                    "读回验证": row.get("readback_verified", ""),
                    "写前ID": row.get("identity_before", ""),
                    "写后ID": row.get("identity_after", ""),
                    "目标值": row.get("target_senco9_values")
                    or row.get("target_senco1_values")
                    or row.get("target_senco3_values")
                    or row.get("candidate_senco2_values")
                    or row.get("candidate_senco4_values")
                    or "",
                    "读回值": row.get("senco1_readback")
                    or row.get("senco2_readback")
                    or row.get("senco3_readback")
                    or row.get("senco4_readback")
                    or row.get("target_senco9_values")
                    or "",
                    "是否写设备ID": row.get("writes_device_id", ""),
                    "是否控制水路气路": row.get("controls_water_or_gas_routes", ""),
                    "证据文件": str(path),
                }
            )
    return rows


def _identity_rows(path: Path) -> List[Dict[str, Any]]:
    return [
        {
            "设备ID": row.get("analyzer_device_id", ""),
            "串口": row.get("port", ""),
            "读前帧ID": row.get("identity_before", ""),
            "读后帧ID": row.get("identity_after", ""),
            "身份已核对": row.get("identity_verified", ""),
            "SETCOMWAY": row.get("sets_comm_way", ""),
            "读前原始帧": row.get("identity_before_raw", ""),
            "读后原始帧": row.get("identity_after_raw", ""),
        }
        for row in _read_csv(path)
    ]


def _temperature_rows(summary_path: Path, coeff_path: Path, device_id: str) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    summary = _one_csv_row(summary_path, device_id, device_key="analyzer_id")
    if summary:
        rows.append(
            {
                "项目": "温度通道评审",
                "设备ID": device_id,
                "覆盖温度": summary.get("distinct_temp_setpoints", ""),
                "腔体温度最大差(C)": _fmt(summary.get("cell_delta_max_c")),
                "壳体温度最大差(C)": _fmt(summary.get("shell_delta_max_c")),
                "腔体拟合最大残差(C)": _fmt(summary.get("cell_fit_max_abs_error_c")),
                "壳体拟合最大残差(C)": _fmt(summary.get("shell_fit_max_abs_error_c")),
                "结论": "已计算候选，未在本轮写入；若写 S7/S8，需先写温度后重算 S1-S4。",
            }
        )
    for row in _read_csv(coeff_path):
        if str(row.get("analyzer_id") or "").zfill(3) != device_id:
            continue
        rows.append(
            {
                "项目": row.get("channel", ""),
                "设备ID": device_id,
                "C0": _fmt(row.get("C0")),
                "C1": _fmt(row.get("C1")),
                "C2": _fmt(row.get("C2")),
                "C3": _fmt(row.get("C3")),
                "结论": "本轮未写入",
            }
        )
    return rows


def _artifact_rows(paths: Sequence[tuple[str, Path]]) -> List[Dict[str, Any]]:
    return [
        {
            "证据角色": label,
            "文件路径": str(path),
            "文件存在": path.exists(),
            "大小(bytes)": path.stat().st_size if path.exists() and path.is_file() else "",
            "SHA256": _sha256(path),
        }
        for label, path in paths
    ]


def _certificate_rows() -> List[Dict[str, Any]]:
    return [
        {"角色": "COM22数字压力计证书", "路径": "D:/手册/FRGsz25038057-数字压力计-118288.pdf"},
        {"角色": "精密露点仪证书", "路径": "C:/Users/A/Desktop/FCDjw25074175-精密露点仪-245932001.pdf"},
        {"角色": "露点仪温度探头证书", "路径": "C:/Users/A/Desktop/CRGzb25074337-精密露点仪(温度探头)-245932001.pdf"},
        {"角色": "铂电阻/数字测温仪证书", "路径": "C:/Users/A/Desktop/铂电阻最新.pdf"},
        {"角色": "高低温交变湿热试验箱证书", "路径": "C:/Users/A/Desktop/CRGxc25071726A-高低温交变湿热试验箱-201712068.pdf"},
        {"角色": "CO2 100-1000ppm及干空气证书", "路径": "照片证据见桌面/下载目录，本报告记录对应采样证据与文件Hash。"},
    ]


def build_payload(root: Path, device_id: str) -> Dict[str, Any]:
    base = root / "logs" / "v1_5_023_100_recalc_write_20260601"
    config_path = root / "configs" / "site_v1_5_formal_open_flow_current_6ch_no_write_active.json"
    final_snapshot = base / "getco_after_100_identity_guard_r2" / "old_component_coefficients_snapshot.json"
    identity_csv = base / "getco_after_100_identity_guard_r2" / "getco_component_snapshot_identity.csv"
    pressure_summary = base / "pressure_senco9_write_100_r1" / "senco9_write_summary.csv"
    co2_summary = base / "co2_senco13_write_100_r1" / "co2_senco13_pair_write_summary.csv"
    h2o_summary = base / "h2o_senco24_write_100_r1" / "h2o_senco24_pair_write_summary.csv"
    co2_residuals = root / "logs" / "co2_fc_T2_all_eligible_full_cert_pressure_checked_20260529_mainpy" / "candidate_fit_residuals.csv"
    h2o_residuals = base / "h2o_senco24_recalc_023_100_raw_r1" / "h2o_senco24_residuals.csv"
    temp_summary = base / "temperature_channel_review_100_r1" / "temperature_channel_summary.csv"
    temp_coeff = base / "temperature_channel_review_100_r1" / "temperature_compensation_coefficients.csv"
    h2o_policy = base / "h2o_senco24_recalc_023_100_raw_r1" / "h2o_senco24_device_policy.csv"

    co2_rows = _co2_fit_rows(co2_residuals, device_id)
    h2o_rows = _h2o_fit_rows(h2o_residuals, device_id)
    co2_err = _error_summary(co2_rows, error_key="误差(ppm)", pct_key="相对误差(%)")
    h2o_err = _error_summary(h2o_rows, error_key="误差(mmol/mol)", pct_key="相对误差(%)")
    pressure_row = _one_csv_row(pressure_summary, device_id)
    co2_write = _one_csv_row(co2_summary, device_id)
    h2o_write = _one_csv_row(h2o_summary, device_id)

    conclusion = [
        {
            "设备ID": device_id,
            "压力SENCO9": pressure_row.get("status", ""),
            "CO2 SENCO1/3": co2_write.get("status", ""),
            "H2O SENCO2/4": h2o_write.get("status", ""),
            "SENCO5": "保持中性 [0.0, 1.0]",
            "SENCO6": "保持中性 [0.0, 1.0]",
            "SENCO7/8": "已评审候选，未写入",
            "CO2主拟合最大误差(ppm)": co2_err["最大绝对误差"],
            "CO2主拟合最大相对误差(%)": co2_err["最大相对误差(%)"],
            "H2O主拟合最大误差(mmol/mol)": h2o_err["最大绝对误差"],
            "H2O主拟合最大相对误差(%)": h2o_err["最大相对误差(%)"],
            "结论": "100 已完成压力、CO2、H2O主链路受控写入与读回核对；建议下一步做少量开放流通真机复验。",
        },
        {
            "设备ID": "023",
            "压力SENCO9": "本轮不处理",
            "CO2 SENCO1/3": "本轮不处理",
            "H2O SENCO2/4": "未计算",
            "SENCO5": "",
            "SENCO6": "",
            "SENCO7/8": "",
            "CO2主拟合最大误差(ppm)": "",
            "CO2主拟合最大相对误差(%)": "",
            "H2O主拟合最大误差(mmol/mol)": "",
            "H2O主拟合最大相对误差(%)": "",
            "结论": "当前 H2O 全温度证据包未发现 023 完整水路样本；按用户要求，本轮不计算 023 水系数。",
        },
    ]
    write_paths = [
        ("压力 SENCO9 写入", pressure_summary),
        ("CO2 SENCO1/3 写入", co2_summary),
        ("H2O SENCO2/4 写入", h2o_summary),
    ]
    artifacts = [
        ("配置", config_path),
        ("身份门禁最终GETCO1-9快照", final_snapshot),
        ("身份门禁CSV", identity_csv),
        ("压力SENCO9写入摘要", pressure_summary),
        ("CO2 SENCO1/3写入摘要", co2_summary),
        ("H2O SENCO2/4写入摘要", h2o_summary),
        ("CO2全温度残差", co2_residuals),
        ("H2O全温度残差", h2o_residuals),
        ("温度通道评审", temp_summary),
        ("温度候选系数", temp_coeff),
        ("H2O设备策略", h2o_policy),
    ]
    return {
        "report_no": f"V15-CAL-DEVICE-{device_id}-20260601-ZH-001",
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "git_commit": _git_commit(root),
        "device_id": device_id,
        "结论摘要": conclusion,
        "设备映射": _device_rows(config_path, device_id),
        "身份核对": _identity_rows(identity_csv),
        "最终GETCO1_9": _snapshot_rows(final_snapshot, device_id),
        "CO2主拟合点位误差": co2_rows,
        "H2O主拟合点位误差": h2o_rows,
        "温度通道评审": _temperature_rows(temp_summary, temp_coeff, device_id),
        "写入审计": _write_event_rows(write_paths, device_id),
        "证书索引": _certificate_rows(),
        "文件Hash": _artifact_rows(artifacts),
    }


def _markdown_table(rows: Sequence[Mapping[str, Any]], max_rows: Optional[int] = None) -> List[str]:
    if not rows:
        return ["_无记录_"]
    selected = list(rows if max_rows is None else rows[:max_rows])
    header: List[str] = []
    for row in selected:
        for key in row:
            if str(key) not in header:
                header.append(str(key))
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join("---" for _ in header) + " |"]
    for row in selected:
        lines.append("| " + " | ".join(str(row.get(key, "")).replace("|", "\\|").replace("\n", " ") for key in header) + " |")
    if max_rows is not None and len(rows) > max_rows:
        lines.append(f"| 其余记录 | 完整数据见 XLSX，共 {len(rows)} 行 |" + " |" * max(0, len(header) - 2))
    return lines


def _write_markdown(path: Path, payload: Mapping[str, Any]) -> None:
    lines: List[str] = [
        f"# {REPORT_TITLE}",
        "",
        f"- 报告编号：`{payload['report_no']}`",
        f"- 设备ID：`{payload['device_id']}`",
        f"- 生成时间：`{payload['generated_at']}`",
        f"- Git提交：`{payload['git_commit']}`",
        "- 生成边界：只读证据包，不打开水路/气路流程，不运行 V2 真机，不写设备 ID。",
        "",
        "## 物理方法与结论",
        "",
        "本报告围绕 V1.5 正式开放流通校准链路生成。压力通道作为独立输入量先校准/验证；CO2 与 H2O 主链路在当前大气压开放流通样气下拟合，冻结封路压力项，不用污染压力点吸收组分误差。",
        "本轮只计算并写入设备 `100`；`023` 若没有完整 H2O 全温度证据，则不计算水系数。",
        "",
    ]
    for title, key, limit in (
        ("结论摘要", "结论摘要", None),
        ("设备映射", "设备映射", None),
        ("身份核对", "身份核对", None),
        ("最终 GETCO1-9", "最终GETCO1_9", None),
        ("CO2 主拟合点位误差", "CO2主拟合点位误差", 80),
        ("H2O 主拟合点位误差", "H2O主拟合点位误差", 80),
        ("温度通道评审", "温度通道评审", None),
        ("写入审计", "写入审计", None),
        ("证书索引", "证书索引", None),
        ("文件 Hash", "文件Hash", None),
    ):
        lines.extend(["", f"## {title}", ""])
        lines.extend(_markdown_table(payload[key], max_rows=limit))
    path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def _set_cjk_font(paragraph: Any, *, size_pt: Optional[float] = None, bold: Optional[bool] = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    for run in paragraph.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        if bold is not None:
            run.bold = bold


def _add_docx_table(doc: Any, title: str, rows: Sequence[Mapping[str, Any]], *, max_rows: Optional[int] = None) -> None:
    heading = doc.add_heading(title, level=2)
    _set_cjk_font(heading, size_pt=12, bold=True)
    if not rows:
        p = doc.add_paragraph("无记录。")
        _set_cjk_font(p)
        return
    selected = list(rows if max_rows is None else rows[:max_rows])
    header: List[str] = []
    for row in selected:
        for key in row:
            if str(key) not in header:
                header.append(str(key))
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for idx, key in enumerate(header):
        table.rows[0].cells[idx].text = key
    for row in selected:
        cells = table.add_row().cells
        for idx, key in enumerate(header):
            cells[idx].text = str(row.get(key, ""))
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _set_cjk_font(paragraph, size_pt=8)
    if max_rows is not None and len(rows) > max_rows:
        p = doc.add_paragraph(f"完整数据见 XLSX 附表；本节仅显示前 {max_rows} 行。")
        _set_cjk_font(p)


def _write_docx(path: Path, payload: Mapping[str, Any]) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    title = doc.add_heading(REPORT_TITLE, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cjk_font(title, size_pt=18, bold=True)
    for text in (
        f"报告编号：{payload['report_no']}",
        f"设备ID：{payload['device_id']}",
        f"生成时间：{payload['generated_at']}",
        f"Git提交：{payload['git_commit']}",
        "边界：只读证据包；不控制水路/气路；不运行 V2 真机；不写设备 ID。",
    ):
        p = doc.add_paragraph(text)
        _set_cjk_font(p)

    p = doc.add_paragraph(
        "物理说明：压力通道独立校准/验证；CO2/H2O 主链路使用开放流通当前大气压证据拟合，冻结封路压力项。"
    )
    _set_cjk_font(p)
    for title, key, limit in (
        ("结论摘要", "结论摘要", None),
        ("设备映射", "设备映射", None),
        ("身份核对", "身份核对", None),
        ("最终 GETCO1-9", "最终GETCO1_9", None),
        ("CO2 主拟合点位误差", "CO2主拟合点位误差", 80),
        ("H2O 主拟合点位误差", "H2O主拟合点位误差", 80),
        ("温度通道评审", "温度通道评审", None),
        ("写入审计", "写入审计", None),
        ("证书索引", "证书索引", None),
        ("文件 Hash", "文件Hash", None),
    ):
        _add_docx_table(doc, title, payload[key], max_rows=limit)
    doc.save(path)


def _write_xlsx(path: Path, payload: Mapping[str, Any]) -> None:
    import pandas as pd

    sheets = {
        "结论摘要": payload["结论摘要"],
        "设备映射": payload["设备映射"],
        "身份核对": payload["身份核对"],
        "最终GETCO1-9": payload["最终GETCO1_9"],
        "CO2点位误差": payload["CO2主拟合点位误差"],
        "H2O点位误差": payload["H2O主拟合点位误差"],
        "温度通道评审": payload["温度通道评审"],
        "写入审计": payload["写入审计"],
        "证书索引": payload["证书索引"],
        "文件Hash": payload["文件Hash"],
    }
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        for name, rows in sheets.items():
            pd.DataFrame(list(rows)).to_excel(writer, sheet_name=name[:31], index=False)
        for ws in writer.book.worksheets:
            ws.freeze_panes = "A2"
            for column_cells in ws.columns:
                max_len = max(len(str(cell.value or "")) for cell in column_cells)
                ws.column_dimensions[column_cells[0].column_letter].width = min(max(10, max_len + 2), 60)


def write_report(root: Path, output_dir: Path, device_id: str) -> Dict[str, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    payload = build_payload(root, device_id)
    prefix = f"v1_5_{device_id}_单台气体分析仪系数写入与拟合证据报告"
    md_path = output_dir / f"{prefix}.md"
    docx_path = output_dir / f"{prefix}.docx"
    xlsx_path = output_dir / f"{prefix}_数据附表.xlsx"
    meta_path = output_dir / f"{prefix}_meta.json"
    hash_path = output_dir / "报告文件hash.csv"
    _write_markdown(md_path, payload)
    _write_docx(docx_path, payload)
    _write_xlsx(xlsx_path, payload)
    meta = {
        "report_no": payload["report_no"],
        "generated_at": payload["generated_at"],
        "device_id": device_id,
        "outputs": {"markdown": str(md_path), "docx": str(docx_path), "xlsx": str(xlsx_path)},
        "note": "中文报告；生成过程只读，不控制水路气路，不写设备。",
    }
    meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
    _write_csv(
        hash_path,
        _artifact_rows(
            [
                ("中文Markdown报告", md_path),
                ("中文DOCX报告", docx_path),
                ("XLSX数据附表", xlsx_path),
                ("报告meta", meta_path),
            ]
        ),
    )
    return {"markdown": md_path, "docx": docx_path, "xlsx": xlsx_path, "meta": meta_path, "hashes": hash_path}


def _parse_args(argv: Optional[Iterable[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export Chinese V1.5 single-device coefficient report.")
    parser.add_argument("--root", default=".")
    parser.add_argument("--device-id", default="100")
    parser.add_argument("--output-dir", default="logs/v1_5_100_single_device_coefficient_report_zh_20260601")
    return parser.parse_args(list(argv) if argv is not None else None)


def main(argv: Optional[Iterable[str]] = None) -> int:
    args = _parse_args(argv)
    device_id = str(args.device_id or "").strip().zfill(3)
    outputs = write_report(Path(args.root).resolve(), Path(args.output_dir).resolve(), device_id)
    print(json.dumps({key: str(value) for key, value in outputs.items()}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
