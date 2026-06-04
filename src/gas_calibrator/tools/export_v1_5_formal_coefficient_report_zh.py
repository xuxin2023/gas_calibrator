"""Export a Chinese V1.5 formal coefficient fitting report.

This report generator is deliberately evidence-only. It reads already-created
V1.5 coefficient snapshots, write-event evidence, and point-error tables. It
does not open serial ports, control gas/water routes, or write coefficients.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import math
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Mapping, Optional, Sequence


REPORT_PREFIX = "v1_5_四通道气体分析仪校准拟合正式报告"


def _now() -> str:
    return datetime.now().isoformat(timespec="seconds")


def _read_csv(path: str | Path) -> List[Dict[str, str]]:
    source = Path(path)
    if not source.exists():
        return []
    with source.open("r", encoding="utf-8-sig", newline="") as handle:
        return [dict(row) for row in csv.DictReader(handle)]


def _write_csv(path: Path, rows: Sequence[Mapping[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    keys: List[str] = []
    for row in rows:
        for key in row:
            if key not in keys:
                keys.append(str(key))
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=keys)
        writer.writeheader()
        writer.writerows(rows)


def _read_json(path: str | Path) -> Any:
    return json.loads(Path(path).read_text(encoding="utf-8"))


def _sha256(path: str | Path) -> str:
    source = Path(path)
    if not source.exists() or not source.is_file():
        return "MISSING"
    digest = hashlib.sha256()
    with source.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest().upper()


def _git_commit(root: Path) -> str:
    try:
        return subprocess.check_output(
            ["git", "rev-parse", "HEAD"],
            cwd=str(root),
            text=True,
            stderr=subprocess.DEVNULL,
        ).strip()
    except Exception:
        return ""


def _fmt(value: Any, digits: int = 6) -> str:
    if value in (None, "", "None", "null"):
        return ""
    try:
        number = float(value)
    except Exception:
        return str(value)
    if not math.isfinite(number):
        return str(value)
    return f"{number:.{digits}g}"


def _json_values(value: Any) -> List[float]:
    if isinstance(value, list):
        return [float(item) for item in value]
    if value in (None, ""):
        return []
    try:
        payload = json.loads(str(value))
    except Exception:
        return []
    if not isinstance(payload, list):
        return []
    out: List[float] = []
    for item in payload:
        try:
            out.append(float(item))
        except Exception:
            pass
    return out


def _snapshot_to_getco_rows(snapshot_path: Path) -> List[Dict[str, Any]]:
    data = _read_json(snapshot_path)
    rows: List[Dict[str, Any]] = []
    for device_id, info in sorted((data.get("analyzers") or data).items()):
        if not isinstance(info, Mapping):
            continue
        for group in range(1, 10):
            values = info.get(f"GETCO{group}_before") or []
            parsed = info.get(f"GETCO{group}_before_parsed") or {}
            if not values and parsed:
                values = [parsed.get(f"C{idx}") for idx in range(6) if f"C{idx}" in parsed]
            values = _json_values(values)
            row: Dict[str, Any] = {
                "设备ID": str(device_id),
                "通道名": info.get("analyzer_prefix", ""),
                "串口": info.get("port", ""),
                "系数组": f"GETCO{group}",
                "读回命令": info.get(f"GETCO{group}_before_command", f"GETCO,YGAS,FFF,{group}"),
                "系数值": json.dumps(values, ensure_ascii=False),
            }
            for idx in range(6):
                row[f"C{idx}"] = _fmt(values[idx]) if idx < len(values) else ""
            rows.append(row)
    return rows


def _load_config_devices(config_path: Path) -> List[Dict[str, Any]]:
    cfg = _read_json(config_path)
    devices = ((cfg.get("devices") or {}).get("gas_analyzers") or [])
    rows: List[Dict[str, Any]] = []
    for item in devices:
        if not bool(item.get("enabled", True)):
            continue
        rows.append(
            {
                "设备ID": item.get("device_id", ""),
                "通道名": item.get("name", ""),
                "串口": item.get("port", ""),
                "波特率": item.get("baud", item.get("baudrate", "")),
                "MODE": item.get("mode", ""),
                "主动上传": item.get("active_send", ""),
                "FTD频率Hz": item.get("ftd_hz", ""),
                "平均/滤波窗口": item.get("average_filter", ""),
                "身份说明": "以气体分析仪自身 MODE2 ID 为身份，不以串口号作为设备身份。",
            }
        )
    return rows


def _co2_final_rows(path: Path) -> List[Dict[str, Any]]:
    rows = []
    for row in _read_csv(path):
        rows.append(
            {
                "设备ID": row.get("device_id", ""),
                "点位": row.get("point_run_id", ""),
                "证书CO2(ppm)": _fmt(row.get("certificate_co2_ppm")),
                "修正前CO2(ppm)": _fmt(row.get("measured_before_final_trim_ppm")),
                "S5_C0": _fmt(row.get("final_S5_C0")),
                "S5_C1": _fmt(row.get("final_S5_C1")),
                "离线反推CO2(ppm)": _fmt(row.get("offline_corrected_co2_ppm")),
                "误差(ppm)": _fmt(row.get("error_ppm")),
                "相对误差(%)": _fmt(row.get("error_pct")),
                "状态": _status_cn(row.get("status", "")),
            }
        )
    return rows


def _h2o_final_rows(path: Path) -> List[Dict[str, Any]]:
    rows = []
    for row in _read_csv(path):
        rows.append(
            {
                "设备ID": row.get("device_id", ""),
                "点位": row.get("point_run_id", ""),
                "参考H2O(mmol/mol)": _fmt(row.get("reference_h2o_mmol")),
                "修正前H2O(mmol/mol)": _fmt(row.get("measured_before_final_trim_mmol")),
                "S6_C0": _fmt(row.get("final_S6_C0")),
                "S6_C1": _fmt(row.get("final_S6_C1")),
                "离线反推H2O(mmol/mol)": _fmt(row.get("offline_corrected_h2o_mmol")),
                "误差(mmol/mol)": _fmt(row.get("error_mmol")),
                "相对误差(%)": _fmt(row.get("error_pct")),
                "状态": _status_cn(row.get("status", "")),
            }
        )
    return rows


def _summary_rows(path: Path, component: str) -> List[Dict[str, Any]]:
    out = []
    unit = "ppm" if component == "CO2" else "mmol/mol"
    for row in _read_csv(path):
        out.append(
            {
                "设备ID": row.get("device_id", ""),
                "点数": row.get("point_count", ""),
                f"最大绝对误差({unit})": _fmt(row.get("max_abs_error")),
                "最大相对误差(%)": _fmt(row.get("max_abs_error_pct")),
            }
        )
    return out


def _raw_rows_cn(rows: Sequence[Mapping[str, Any]], limit: Optional[int] = None) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for idx, row in enumerate(rows):
        if limit is not None and idx >= limit:
            break
        out.append({str(key): value for key, value in row.items()})
    return out


def _status_cn(value: Any) -> str:
    text = str(value or "").strip()
    mapping = {
        "written_or_retained": "已写入或已保留",
        "not_written_blocked_final_trim": "最终修正候选被阻断，未写入",
        "written_safe_nonnegative_c0": "按安全可读回值写入",
        "written_readback_verified": "已写入并读回确认",
        "written_readback_verified_ack_missing": "ACK缺失但读回匹配",
        "failed": "失败",
        "blocked": "阻断",
        "already_neutral": "已为中性",
    }
    return mapping.get(text, text)


def _write_events(paths: Sequence[tuple[str, Path]]) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for source, path in paths:
        for row in _read_csv(path):
            rows.append(
                {
                    "事件来源": source,
                    "设备ID": row.get("device_id", ""),
                    "串口": row.get("port", ""),
                    "旧值": row.get("old_senco5") or row.get("old_senco6") or "",
                    "目标值": row.get("target_senco5") or row.get("target_senco6") or "",
                    "最终读回": row.get("final_senco5") or row.get("final_senco6") or "",
                    "命令": row.get("payload", ""),
                    "ACK/响应": row.get("write_ack_response_json", ""),
                    "状态": _status_cn(row.get("status", "")),
                    "原因": row.get("reason", ""),
                    "是否控制水路/气路": row.get("controls_water_or_gas_routes", ""),
                    "证据文件": str(path),
                }
            )
    return rows


def _certificate_sources(root: Path) -> List[Dict[str, Any]]:
    def item(role: str, value: str, uncertainty: str, path: str) -> Dict[str, Any]:
        p = Path(path)
        return {
            "证据角色": role,
            "证书值/说明": value,
            "不确定度": uncertainty,
            "文件路径": str(p),
            "文件存在": p.exists(),
            "SHA256": _sha256(p),
        }

    desktop = Path("C:/Users/A/Desktop")
    return [
        item("CO2 标气 100 ppm", "99.94 umol/mol", "1%", str(desktop / "微信图片_20260527154435_237_36.jpg")),
        item("CO2 标气 200 ppm", "200.10 umol/mol", "1%", str(desktop / "微信图片_20260528161418_246_36.jpg")),
        item("CO2 标气 300 ppm", "299.73 umol/mol", "1%", str(desktop / "微信图片_20260527160742_240_36.jpg")),
        item("CO2 标气 400 ppm", "399.56 umol/mol", "1%", str(desktop / "微信图片_20260528161507_248_36.jpg")),
        item("CO2 标气 500 ppm", "500.13 umol/mol", "2%", str(desktop / "微信图片_20260527160740_238_36.jpg")),
        item("CO2 标气 600 ppm", "599.54 umol/mol", "1%", str(desktop / "微信图片_20260528161556_250_36.jpg")),
        item("CO2 标气 700 ppm", "700.33 umol/mol", "1%", str(desktop / "微信图片_20260527161922_242_36.jpg")),
        item("CO2 标气 800 ppm", "800.59 umol/mol", "1%", str(desktop / "微信图片_20260528161648_252_36.jpg")),
        item("CO2 标气 900 ppm", "897.04 umol/mol", "1%", str(desktop / "微信图片_20260525191354_234_36.jpg")),
        item("CO2 标气 1000 ppm", "998.62 umol/mol", "1%", str(desktop / "微信图片_20260528161731_254_36.jpg")),
        item("干空气/低端锚点", "O2 20.95%, N2 平衡；CO2 仅作低端工程锚点", "不是 CO2 零气证书", str(desktop / "微信图片_20260528161327_244_36.jpg")),
        item("COM22 数字压力计证书", "压力参考", "见证书", "D:/手册/FRGsz25038057-数字压力计-118288.pdf"),
        item("精密露点仪证书", "露点/湿度参考", "见证书", str(desktop / "FCDjw25074175-精密露点仪-245932001.pdf")),
        item("露点仪温度探头证书", "温度参考", "见证书", str(desktop / "CRGzb25074337-精密露点仪(温度探头)-245932001.pdf")),
        item("铂电阻数字测温仪证书", "温箱/腔体温度参考", "见证书", str(desktop / "铂电阻最新.pdf")),
        item("高低温交变湿热试验箱证书", "温度环境设备", "见证书", str(desktop / "CRGxc25071726A-高低温交变湿热试验箱-201712068.pdf")),
    ]


def _artifact_hashes(paths: Sequence[tuple[str, Path]]) -> List[Dict[str, Any]]:
    rows = []
    for role, path in paths:
        rows.append(
            {
                "证据角色": role,
                "文件路径": str(path),
                "文件存在": path.exists(),
                "大小(bytes)": path.stat().st_size if path.exists() and path.is_file() else "",
                "SHA256": _sha256(path),
            }
        )
    return rows


def _markdown_table(rows: Sequence[Mapping[str, Any]], max_rows: Optional[int] = None) -> List[str]:
    if not rows:
        return ["_无记录_"]
    selected = list(rows if max_rows is None else rows[:max_rows])
    header: List[str] = []
    for row in selected:
        for key in row:
            text = str(key)
            if text not in header:
                header.append(text)
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join("---" for _ in header) + " |"]
    for row in selected:
        lines.append("| " + " | ".join(str(row.get(key, "")).replace("|", "\\|").replace("\n", " ") for key in header) + " |")
    if max_rows is not None and len(rows) > max_rows:
        lines.append(f"| 其余记录 | 完整数据见 XLSX，共 {len(rows)} 行 |" + " |" * max(0, len(header) - 2))
    return lines


def _write_markdown(path: Path, payload: Mapping[str, Any]) -> None:
    lines: List[str] = [
        "# V1.5 四通道气体分析仪校准拟合正式报告",
        "",
        f"- 报告编号：`{payload['report_no']}`",
        f"- 生成时间：`{payload['generated_at']}`",
        f"- Git 提交：`{payload['git_commit']}`",
        "- 报告语言：中文",
        "",
        "## 结论摘要",
        "",
        "本报告从已冻结的 V1.5 开放流通全温度采样证据、系数写入事件和最终 GETCO1-9 读回快照生成。",
        "本次没有重新打开水路、气路或压力控制流程；报告生成过程不写设备、不改设备 ID、不运行 V2 真机。",
        "",
    ]
    lines.extend(_markdown_table(payload["结论摘要"]))
    lines.extend(
        [
            "",
            "## 物理方法与边界",
            "",
            "- CO2/H2O 主校准基于开放流通样气，标准气或湿气持续刷新分析仪腔体和下游管路。",
            "- 压力通道已作为独立输入量处理；当前 CO2/H2O 拟合不引入封路多压力点，也不把压力污染点作为正式拟合输入。",
            "- SENCO1/SENCO3 为 CO2 主光学/温度链路，SENCO2/SENCO4 为 H2O 主光学/温度链路。",
            "- SENCO5/SENCO6 是最终显示浓度线性修正层，物理模型为：最终值 = 主链路输出 * C1 + C0。",
            "- SENCO5/SENCO6 写入允许负值，payload 小数位最多 3 位；已验证不可靠的设备/参数组合必须阻断或保留安全值。",
            "- 033 的 SENCO6 理论最优负小数 `[-0.1, 1.1]` 已被单独审计；该设备固件对负小数写入/读回不可靠，因此最终保留可读回安全值 `[0.0, 1.1]`。",
            "- 本报告的“最终修正点位误差”使用已完成的真实复测采样点作为输入，再带入最终 GETCO 读回的 S5/S6 线性修正离线反推；它不是重新写入后再次在线采样得到的直接显示误差。",
            "",
            "## 设备参数",
            "",
        ]
    )
    lines.extend(_markdown_table(payload["设备参数"]))
    lines.extend(["", "## 最终 GETCO1-9 系数", ""])
    lines.extend(_markdown_table(payload["最终GETCO1_9"], max_rows=60))
    lines.extend(["", "## CO2 最终修正点位误差", ""])
    lines.extend(_markdown_table(payload["CO2最终误差"]))
    lines.extend(["", "## H2O 最终修正点位误差", ""])
    lines.extend(_markdown_table(payload["H2O最终误差"]))
    lines.extend(["", "## 全温度主拟合误差摘要", ""])
    lines.extend(_markdown_table(payload["主拟合摘要"]))
    lines.extend(["", "## 写入审计", ""])
    lines.extend(_markdown_table(payload["写入审计"], max_rows=80))
    lines.extend(["", "## 溯源证书与参考设备", ""])
    lines.extend(_markdown_table(payload["溯源证书"], max_rows=80))
    lines.extend(["", "## 文件 Hash", ""])
    lines.extend(_markdown_table(payload["文件Hash"], max_rows=120))
    path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def _set_run_font(run: Any, *, ascii_font: str = "Calibri", east_asia_font: str = "宋体", size_pt: Optional[float] = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def _set_paragraph_font(paragraph: Any, *, east_asia_font: str = "宋体", size_pt: Optional[float] = None) -> None:
    for run in paragraph.runs:
        _set_run_font(run, east_asia_font=east_asia_font, size_pt=size_pt)


def _add_docx_table(doc: Any, title: str, rows: Sequence[Mapping[str, Any]], *, max_rows: Optional[int] = None) -> None:
    heading = doc.add_heading(title, level=2)
    _set_paragraph_font(heading, east_asia_font="黑体", size_pt=12)
    if not rows:
        p = doc.add_paragraph("无记录。")
        _set_paragraph_font(p)
        return
    selected = list(rows if max_rows is None else rows[:max_rows])
    header: List[str] = []
    for row in selected:
        for key in row:
            text = str(key)
            if text not in header:
                header.append(text)
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for idx, key in enumerate(header):
        cell = table.rows[0].cells[idx]
        cell.text = key
    for row in selected:
        cells = table.add_row().cells
        for idx, key in enumerate(header):
            cells[idx].text = str(row.get(key, ""))
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _set_paragraph_font(paragraph, size_pt=8.5)
    if max_rows is not None and len(rows) > max_rows:
        p = doc.add_paragraph(f"完整数据见同目录 XLSX 附表，本节仅显示前 {max_rows} 行；总记录数 {len(rows)} 行。")
        _set_paragraph_font(p)


def _write_docx(path: Path, payload: Mapping[str, Any]) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    for style_name in ("Normal", "Body Text"):
        try:
            style = styles[style_name]
            style.font.name = "Calibri"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            style.font.size = Pt(10.5)
        except Exception:
            pass

    title = doc.add_heading("V1.5 四通道气体分析仪校准拟合正式报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_font(title, east_asia_font="黑体", size_pt=18)
    for text in (
        f"报告编号：{payload['report_no']}",
        f"生成时间：{payload['generated_at']}",
        f"Git 提交：{payload['git_commit']}",
        "报告生成过程：只读证据包，不控制气路/水路/压力控制器，不写设备。",
    ):
        p = doc.add_paragraph(text)
        _set_paragraph_font(p)

    h = doc.add_heading("一、结论摘要", level=1)
    _set_paragraph_font(h, east_asia_font="黑体", size_pt=14)
    for text in (
        "本报告依据 V1.5 开放流通校准证据、受控写入事件和最终 GETCO1-9 读回快照生成。",
        "CO2/H2O 主链路拟合不引入封路压力项；压力通道作为独立输入量验证，不用污染压力点吸收组分误差。",
        "033 的 SENCO6 负小数理论目标不作为最终写入状态，最终报告以实际可读回值为准。",
    ):
        p = doc.add_paragraph(text)
        _set_paragraph_font(p)
    _add_docx_table(doc, "结论摘要", payload["结论摘要"])

    h = doc.add_heading("二、物理方法与计量边界", level=1)
    _set_paragraph_font(h, east_asia_font="黑体", size_pt=14)
    for text in (
        "开放流通的物理意义是持续刷新分析仪光学腔体和管路，使分析仪看到稳定、可追溯、未被封路死体积污染的标准气或湿气。",
        "SENCO1/SENCO3 与 SENCO2/SENCO4 负责主光学/温度链路，SENCO5/SENCO6 只作为最终显示浓度线性层：最终值 = 主链路输出 * C1 + C0。",
        "相对误差用于评价非零工作点的读数比例偏差；绝对误差用于评价零点、低端点和截距。二者同时保留，不能互相替代。",
    ):
        p = doc.add_paragraph(text)
        _set_paragraph_font(p)

    _add_docx_table(doc, "设备参数", payload["设备参数"])
    _add_docx_table(doc, "最终 GETCO1-9 系数", payload["最终GETCO1_9"], max_rows=60)
    _add_docx_table(doc, "CO2 最终修正点位误差", payload["CO2最终误差"])
    _add_docx_table(doc, "H2O 最终修正点位误差", payload["H2O最终误差"])
    _add_docx_table(doc, "全温度主拟合误差摘要", payload["主拟合摘要"])
    _add_docx_table(doc, "写入审计", payload["写入审计"], max_rows=80)
    _add_docx_table(doc, "溯源证书与参考设备", payload["溯源证书"], max_rows=80)
    _add_docx_table(doc, "文件 Hash", payload["文件Hash"], max_rows=120)

    doc.save(path)


def _write_xlsx(path: Path, payload: Mapping[str, Any]) -> None:
    import pandas as pd

    sheet_map = {
        "结论摘要": payload["结论摘要"],
        "设备参数": payload["设备参数"],
        "最终GETCO1-9": payload["最终GETCO1_9"],
        "CO2最终误差": payload["CO2最终误差"],
        "H2O最终误差": payload["H2O最终误差"],
        "CO2主拟合摘要": payload["CO2主拟合摘要"],
        "H2O主拟合摘要": payload["H2O主拟合摘要"],
        "CO2全温误差": payload["CO2全温误差"],
        "H2O全温误差": payload["H2O全温误差"],
        "写入审计": payload["写入审计"],
        "溯源证书": payload["溯源证书"],
        "文件Hash": payload["文件Hash"],
    }
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        for name, rows in sheet_map.items():
            pd.DataFrame(list(rows)).to_excel(writer, sheet_name=name[:31], index=False)
        for ws in writer.book.worksheets:
            ws.freeze_panes = "A2"
            for column_cells in ws.columns:
                max_len = max(len(str(cell.value or "")) for cell in column_cells)
                ws.column_dimensions[column_cells[0].column_letter].width = min(max(10, max_len + 2), 48)


def build_payload(root: Path) -> Dict[str, Any]:
    output_source = root / "logs" / "v1_5_formal_coefficient_report_20260601"
    final_snapshot = root / "logs" / "v1_5_final_linear_trim_write_20260601" / "readonly_final_getco1_9_snapshot_after_report_fix_20260601_r3" / "old_component_coefficients_snapshot.json"
    if not final_snapshot.exists():
        final_snapshot = root / "logs" / "v1_5_final_linear_trim_write_20260601" / "readonly_final_getco1_9_snapshot_20260601_r2" / "old_component_coefficients_snapshot.json"
    config_path = root / "configs" / "site_v1_5_formal_open_flow_current_4ch_no_write_active_generated_20260531.json"
    co2_final = output_source / "co2_final_trim_point_errors.csv"
    h2o_final = output_source / "h2o_final_trim_point_errors.csv"
    co2_final_summary = output_source / "co2_final_trim_summary.csv"
    h2o_final_summary = output_source / "h2o_final_trim_summary.csv"
    co2_main_summary = output_source / "co2_main_fit_summary.csv"
    h2o_main_summary = output_source / "h2o_main_fit_summary.csv"
    co2_full = output_source / "co2_fulltemp_main_fit_point_errors.csv"
    h2o_full = output_source / "h2o_fulltemp_main_fit_point_errors.csv"

    co2_summary = _summary_rows(co2_final_summary, "CO2")
    h2o_summary = _summary_rows(h2o_final_summary, "H2O")
    by_device: Dict[str, Dict[str, Any]] = {}
    for row in co2_summary:
        by_device.setdefault(str(row["设备ID"]), {})["CO2最大相对误差(%)"] = row["最大相对误差(%)"]
        by_device[str(row["设备ID"])]["CO2最大绝对误差(ppm)"] = row["最大绝对误差(ppm)"]
    for row in h2o_summary:
        by_device.setdefault(str(row["设备ID"]), {})["H2O最大相对误差(%)"] = row["最大相对误差(%)"]
        by_device[str(row["设备ID"])]["H2O最大绝对误差(mmol/mol)"] = row["最大绝对误差(mmol/mol)"]
    conclusion = []
    for dev in sorted(by_device):
        status = "可进入小范围真机复验"
        note = "最终状态以 GETCO 读回和离线反推为准。"
        if dev == "030":
            status = "主链路保留，S5/S6最终修正未写入"
            note = "030 的最终线性修正幅度过大，当前报告保留中性 S5/S6。"
        if dev == "033":
            note = "033 的 S6 理论负小数目标不可靠，最终采用可读回安全值 [0.0, 1.1]。"
        conclusion.append({"设备ID": dev, **by_device[dev], "状态": status, "说明": note})

    write_paths = [
        ("S5 初次写入失败审计", root / "logs" / "v1_5_final_linear_trim_write_20260601" / "senco5_pctopt_skip030_022_033_051_20260601_r2" / "senco5_linear_write_events.csv"),
        ("S5 清除到中性", root / "logs" / "v1_5_final_linear_trim_write_20260601" / "senco5_clear_022_033_051_before_rewrite_20260601_r3" / "senco5_neutral_write_events.csv"),
        ("S5 最终写入", root / "logs" / "v1_5_final_linear_trim_write_20260601" / "senco5_pctopt_after_clear_skip030_022_033_051_20260601_r4" / "senco5_linear_write_events.csv"),
        ("033 S6 清除到中性", root / "logs" / "v1_5_final_linear_trim_write_20260601" / "senco6_clear_033_before_negative_rewrite_20260601_r5" / "senco6_neutral_write_events.csv"),
        ("033 S6 负小数尝试", root / "logs" / "v1_5_final_linear_trim_write_20260601" / "senco6_negative_033_after_clear_20260601_r6" / "senco6_linear_write_events.csv"),
        ("033 S6 恢复安全值", root / "logs" / "v1_5_final_linear_trim_write_20260601" / "senco6_restore_033_safe_after_negative_attempt_20260601_r7" / "senco6_linear_write_events.csv"),
    ]
    evidence_paths = [
        ("配置快照", config_path),
        ("最终 GETCO1-9 快照", final_snapshot),
        ("CO2 最终修正点位误差", co2_final),
        ("H2O 最终修正点位误差", h2o_final),
        ("CO2 全温度主拟合点位误差", co2_full),
        ("H2O 全温度主拟合点位误差", h2o_full),
        *write_paths,
    ]
    return {
        "report_no": "V15-CAL-FIT-20260601-4CH-ZH-002",
        "generated_at": _now(),
        "git_commit": _git_commit(root),
        "结论摘要": conclusion,
        "设备参数": _load_config_devices(config_path),
        "最终GETCO1_9": _snapshot_to_getco_rows(final_snapshot),
        "CO2最终误差": _co2_final_rows(co2_final),
        "H2O最终误差": _h2o_final_rows(h2o_final),
        "CO2主拟合摘要": _raw_rows_cn(_read_csv(co2_main_summary)),
        "H2O主拟合摘要": _raw_rows_cn(_read_csv(h2o_main_summary)),
        "主拟合摘要": [
            {"组分": "CO2", **row} for row in _raw_rows_cn(_read_csv(co2_main_summary))
        ]
        + [{"组分": "H2O", **row} for row in _raw_rows_cn(_read_csv(h2o_main_summary))],
        "CO2全温误差": _raw_rows_cn(_read_csv(co2_full)),
        "H2O全温误差": _raw_rows_cn(_read_csv(h2o_full)),
        "写入审计": _write_events(write_paths),
        "溯源证书": _certificate_sources(root),
        "文件Hash": _artifact_hashes(evidence_paths),
    }


def write_report(root: Path, output_dir: Path) -> Dict[str, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    payload = build_payload(root)
    md_path = output_dir / f"{REPORT_PREFIX}.md"
    docx_path = output_dir / f"{REPORT_PREFIX}.docx"
    xlsx_path = output_dir / f"{REPORT_PREFIX}_数据附表.xlsx"
    meta_path = output_dir / f"{REPORT_PREFIX}_meta.json"
    hash_path = output_dir / "报告文件hash.csv"

    _write_markdown(md_path, payload)
    _write_docx(docx_path, payload)
    _write_xlsx(xlsx_path, payload)
    payload_meta = {
        "report_no": payload["report_no"],
        "generated_at": payload["generated_at"],
        "git_commit": payload["git_commit"],
        "outputs": {
            "markdown": str(md_path),
            "docx": str(docx_path),
            "xlsx": str(xlsx_path),
        },
        "note": "中文正式报告；DOCX 显式设置中文字体；生成过程只读证据，不控制气路/水路，不写设备。",
    }
    meta_path.write_text(json.dumps(payload_meta, ensure_ascii=False, indent=2), encoding="utf-8")
    hashes = _artifact_hashes(
        [
            ("中文正式报告 Markdown", md_path),
            ("中文正式报告 DOCX", docx_path),
            ("中文正式报告 XLSX 附表", xlsx_path),
            ("中文正式报告 meta", meta_path),
        ]
    )
    _write_csv(hash_path, hashes)
    return {"markdown": md_path, "docx": docx_path, "xlsx": xlsx_path, "meta": meta_path, "hashes": hash_path}


def _parse_args(argv: Optional[Iterable[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export Chinese V1.5 formal coefficient fitting report.")
    parser.add_argument("--root", default=".", help="Repository/worktree root.")
    parser.add_argument(
        "--output-dir",
        default="logs/v1_5_formal_coefficient_report_zh_20260601",
        help="Output directory.",
    )
    return parser.parse_args(list(argv) if argv is not None else None)


def main(argv: Optional[Iterable[str]] = None) -> int:
    args = _parse_args(argv)
    root = Path(args.root).resolve()
    outputs = write_report(root, Path(args.output_dir).resolve())
    print(json.dumps({key: str(value) for key, value in outputs.items()}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
