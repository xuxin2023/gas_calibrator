"""Generate V1.5 calibration reports from the evidence package.

The report generator is sidecar-only. It reads evidence artifacts and writes
Markdown/DOCX/PDF files; it never opens COM ports, controls routes or valves,
or writes analyzer coefficients.
"""

from __future__ import annotations

import csv
import hashlib
import json
import math
import statistics
from dataclasses import asdict, dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Mapping, Optional, Sequence


@dataclass
class ReportTable:
    title: str
    rows: List[Dict[str, Any]] = field(default_factory=list)


@dataclass
class ReportSection:
    title: str
    paragraphs: List[str] = field(default_factory=list)
    tables: List[ReportTable] = field(default_factory=list)


@dataclass
class ReportDocument:
    title: str
    sections: List[ReportSection] = field(default_factory=list)


def _now() -> str:
    return datetime.now().isoformat(timespec="seconds")


def _load_json(path: str | Path) -> Dict[str, Any]:
    return json.loads(Path(path).read_text(encoding="utf-8"))


def _load_csv(path: str | Path) -> List[Dict[str, Any]]:
    with Path(path).open("r", encoding="utf-8-sig", newline="") as handle:
        return [dict(row) for row in csv.DictReader(handle)]


def _safe_float(value: Any) -> Optional[float]:
    if value in (None, ""):
        return None
    try:
        numeric = float(value)
    except Exception:
        return None
    if not math.isfinite(numeric):
        return None
    return numeric


def _fmt(value: Any, digits: int = 3) -> str:
    if value in (None, ""):
        return "not_evaluated"
    if isinstance(value, float):
        return f"{value:.{digits}f}"
    return str(value)


def _bool_value(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    return str(value or "").strip().lower() in {"true", "1", "yes", "pass", "ok"}


def _split_reasons(value: Any) -> List[str]:
    if value in (None, ""):
        return []
    if isinstance(value, list):
        return [str(item) for item in value if str(item)]
    return [item for item in str(value).split(";") if item]


def _first_value(row: Mapping[str, Any], keys: Iterable[str]) -> Any:
    for key in keys:
        value = row.get(key)
        if value not in (None, ""):
            return value
    return None


def _first_float(row: Mapping[str, Any], keys: Iterable[str]) -> Optional[float]:
    return _safe_float(_first_value(row, keys))


def _artifact_paths_by_name(bundle: Mapping[str, Any]) -> Dict[str, Path]:
    tables = bundle.get("tables") if isinstance(bundle.get("tables"), Mapping) else {}
    rows = tables.get("sample_files") or []
    out: Dict[str, Path] = {}
    for row in rows:
        if not isinstance(row, Mapping):
            continue
        path_text = str(row.get("path") or "")
        if not path_text:
            continue
        path = Path(path_text)
        out.setdefault(path.name, path)
    return out


def _artifact_rows(bundle: Mapping[str, Any], filename: str) -> List[Dict[str, Any]]:
    path = _artifact_paths_by_name(bundle).get(filename)
    if path is None or not path.exists():
        return []
    try:
        return _load_csv(path)
    except Exception:
        return []


def _artifact_rows_all(bundle: Mapping[str, Any], filename: str) -> List[Dict[str, Any]]:
    tables = bundle.get("tables") if isinstance(bundle.get("tables"), Mapping) else {}
    rows = tables.get("sample_files") or []
    out: List[Dict[str, Any]] = []
    for row in rows:
        if not isinstance(row, Mapping):
            continue
        path_text = str(row.get("path") or "")
        if not path_text:
            continue
        path = Path(path_text)
        if path.name != filename or not path.exists():
            continue
        try:
            out.extend(_load_csv(path))
        except Exception:
            continue
    return out


def _artifact_json(bundle: Mapping[str, Any], filename: str) -> Dict[str, Any]:
    path = _artifact_paths_by_name(bundle).get(filename)
    if path is None or not path.exists():
        return {}
    try:
        payload = _load_json(path)
    except Exception:
        return {}
    return payload if isinstance(payload, Mapping) else {}


def _artifact_manifest_paths(bundle: Mapping[str, Any], role_prefix: str) -> List[str]:
    tables = bundle.get("tables") if isinstance(bundle.get("tables"), Mapping) else {}
    rows = tables.get("sample_files") or []
    paths: List[str] = []
    for row in rows:
        if not isinstance(row, Mapping):
            continue
        role = str(row.get("artifact_role") or "")
        if role.startswith(role_prefix):
            paths.append(str(row.get("path") or ""))
    return [path for path in paths if path]


def _post_write_reverification_model(bundle: Mapping[str, Any]) -> Dict[str, Any]:
    review = _artifact_json(bundle, "post_write_reverification_review.json")
    device_summary = _artifact_rows(bundle, "post_write_reverification_device_summary.csv")
    point_results = _artifact_rows(bundle, "post_write_reverification_points.csv")
    status = str(review.get("overall_status") or "").strip()
    if not status and device_summary:
        statuses = {str(row.get("status") or "").strip().lower() for row in device_summary}
        if statuses and statuses <= {"pass"}:
            status = "pass"
        elif "fail" in statuses:
            status = "fail"
        else:
            status = "review_required"
    return {
        "available": bool(review or device_summary or point_results),
        "overall_status": status or "not_available",
        "created_at": review.get("created_at") or "",
        "limits": review.get("limits") or {},
        "warnings": review.get("warnings") or [],
        "device_summary": device_summary,
        "point_results": point_results[:500],
        "source_artifacts": _artifact_manifest_paths(bundle, "post_write_reverification"),
    }


def _run_evidence_status_model(bundle: Mapping[str, Any]) -> Dict[str, Any]:
    status = _artifact_json(bundle, "v1_5_run_evidence_status.json")
    source_artifacts = _artifact_manifest_paths(bundle, "run_evidence_status")
    return {
        "available": bool(status),
        "overall_status": status.get("overall_status"),
        "current_stage": status.get("current_stage"),
        "contract_status": status.get("contract_status"),
        "stage_statuses": list(status.get("stage_statuses") or []),
        "physical_boundaries": dict(status.get("physical_boundaries") or {}),
        "source_artifacts": source_artifacts,
    }


def _table_rows(bundle: Mapping[str, Any], name: str) -> List[Dict[str, Any]]:
    tables = bundle.get("tables") if isinstance(bundle.get("tables"), Mapping) else {}
    rows = tables.get(name) if isinstance(tables, Mapping) else []
    return [dict(row) for row in rows or [] if isinstance(row, Mapping)]


def _mean(values: Sequence[float]) -> Optional[float]:
    return float(statistics.fmean(values)) if values else None


def _stdev(values: Sequence[float]) -> Optional[float]:
    return float(statistics.stdev(values)) if len(values) > 1 else None


def _component_standard_value(standard_gases: Sequence[Mapping[str, Any]], component: str) -> Optional[float]:
    for gas in standard_gases:
        if str(gas.get("component") or "").strip().lower() == component:
            return _safe_float(gas.get("certificate_value"))
    return None


def _component_uncertainty(standard_gases: Sequence[Mapping[str, Any]], component: str) -> Optional[float]:
    for gas in standard_gases:
        if str(gas.get("component") or "").strip().lower() == component:
            return _safe_float(gas.get("certificate_uncertainty"))
    return None


def _group_component_results(
    rows: Sequence[Mapping[str, Any]],
    standard_gases: Sequence[Mapping[str, Any]],
    *,
    analyzer_prefix: str,
) -> List[Dict[str, Any]]:
    grouped: Dict[str, List[Mapping[str, Any]]] = {}
    for row in rows:
        component = str(row.get("component") or row.get("point_phase") or row.get("route") or "").strip().lower()
        if component not in {"co2", "h2o"}:
            continue
        point = str(row.get("point_row") or row.get("point_tag") or row.get("point_phase") or component)
        grouped.setdefault(f"{component}:{point}", []).append(row)

    result_rows: List[Dict[str, Any]] = []
    for key, items in sorted(grouped.items()):
        component, point = key.split(":", 1)
        if component == "h2o":
            values = [
                value
                for value in (
                    _first_float(
                        row,
                        (
                            f"{analyzer_prefix}_h2o_mmol",
                            "h2o_mmol",
                            "h2o_mmol_mol",
                        ),
                    )
                    for row in items
                )
                if value is not None
            ]
            unit = "mmol/mol"
        else:
            values = [
                value
                for value in (
                    _first_float(
                        row,
                        (
                            f"{analyzer_prefix}_co2_ppm",
                            "co2_ppm",
                            "sample_co2_ppm",
                        ),
                    )
                    for row in items
                )
                if value is not None
            ]
            unit = "ppm"
        standard = _component_standard_value(standard_gases, component)
        measured = _mean(values)
        stdev = _stdev(values)
        result_rows.append(
            {
                "component": component.upper(),
                "point_id": point,
                "standard_value": _fmt(standard),
                "measured_mean": _fmt(measured),
                "error": _fmt(None if standard is None or measured is None else measured - standard),
                "std_dev": _fmt(stdev),
                "sample_count": len(values),
                "unit": unit,
                "qc_grade": "A",
                "entered_formal_fit": True,
                "expanded_uncertainty_k2": "not_released",
            }
        )
    return result_rows


def _device_id_from_row(row: Mapping[str, Any]) -> str:
    for key in ("analyzer_device_id", "device_id", "analyzer_id", "sensor_id"):
        value = row.get(key)
        if value not in (None, ""):
            return str(value)
    return ""


def _is_single_report_analyzer_id(device_id: Any) -> bool:
    text = str(device_id or "").strip()
    if not text:
        return False
    lower = text.lower()
    if lower in {"all", "unknown", "not_available", "not_specified", "n/a"}:
        return False
    if any(marker in text for marker in (";", ",", "|", "[", "]", "{", "}")):
        return False
    if len(text.split()) != 1:
        return False
    return True


def _open_flow_fit_allowed(row: Mapping[str, Any]) -> bool:
    return _bool_value(row.get("candidate_fit_allowed"))


def _point_identity(row: Mapping[str, Any]) -> str:
    return str(
        _first_value(
            row,
            (
                "point_id",
                "point_run_id",
                "point_tag",
                "point_key",
                "point_row",
                "point_phase",
            ),
        )
        or ""
    ).strip()


def _append_reason(existing: Any, reason: str) -> str:
    parts = [item for item in str(existing or "").split(";") if item]
    if reason and reason not in parts:
        parts.append(reason)
    return ";".join(parts)


def _h2o_queue_exclusion_rows(bundle: Mapping[str, Any]) -> List[Dict[str, Any]]:
    rows = _artifact_rows_all(bundle, "queue_abort_exclusion.csv")
    if rows:
        return rows
    payload = _artifact_json(bundle, "queue_abort_exclusion.json")
    raw_rows = payload.get("rows") if isinstance(payload, Mapping) else []
    return [dict(row) for row in raw_rows or [] if isinstance(row, Mapping)]


def _apply_h2o_queue_exclusions(
    open_flow_summary: Sequence[Mapping[str, Any]],
    exclusion_rows: Sequence[Mapping[str, Any]],
) -> List[Dict[str, Any]]:
    excluded_point_ids = {
        str(row.get("point_id") or row.get("point_run_id") or "").strip()
        for row in exclusion_rows
        if _bool_value(row.get("exclude_from_fit"))
        or _bool_value(row.get("exclude_from_acceptance"))
        or _bool_value(row.get("exclude_from_senco_review"))
    }
    excluded_point_ids.discard("")
    if not excluded_point_ids:
        return [dict(row) for row in open_flow_summary]

    out: List[Dict[str, Any]] = []
    for row in open_flow_summary:
        item = dict(row)
        component = str(item.get("component") or "").strip().lower()
        if component == "h2o" and _point_identity(item) in excluded_point_ids:
            item["point_calibratability_grade"] = "C"
            item["point_calibratability_role"] = "diagnostic_only_queue_exclusion"
            item["candidate_fit_allowed"] = False
            item["candidate_fit_blockers"] = _append_reason(
                item.get("candidate_fit_blockers"),
                "h2o_queue_abort_exclusion",
            )
            item["sample_readiness_status"] = "excluded_from_formal_fit"
            item["sample_readiness_blockers"] = _append_reason(
                item.get("sample_readiness_blockers"),
                "h2o_queue_abort_exclusion",
            )
            item["queue_exclusion_status"] = "excluded"
        out.append(item)
    return out


def _sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _certificate_artifact_metadata(
    key: str,
    path: Path,
) -> Dict[str, Any] | None:
    if not key.startswith("device_"):
        return None
    marker = "_calibration_certificate_"
    certificate_kind = "calibration"
    if marker not in key:
        marker = "_verification_certificate_"
        certificate_kind = "verification"
    if marker not in key:
        return None
    device_id = key[len("device_") : key.index(marker)]
    stat = path.stat()
    return {
        "artifact_key": key,
        "device_id": device_id,
        "certificate_kind": certificate_kind,
        "format": path.suffix.lower().lstrip(".") or "unknown",
        "path": str(path),
        "sha256": _sha256_file(path),
        "size_bytes": int(stat.st_size),
    }


def _write_per_device_certificate_manifest(
    *,
    outputs: Mapping[str, Path],
    output_dir: Path,
    model: Mapping[str, Any],
) -> Dict[str, Path]:
    rows: List[Dict[str, Any]] = []
    for key, path in sorted(outputs.items()):
        metadata = _certificate_artifact_metadata(key, path)
        if metadata:
            rows.append(metadata)

    manifest_payload = {
        "schema": "v1_5_per_device_certificate_manifest_v1",
        "generated_at": _now(),
        "run_id": str(model.get("run_id") or ""),
        "report_no": str(model.get("report_no") or ""),
        "physical_boundaries": {
            "offline_report_pack_only": True,
            "opens_com_ports": False,
            "controls_water_or_gas_routes": False,
            "controls_valves_or_pace": False,
            "writes_coefficients": False,
            "not_real_acceptance_evidence": True,
        },
        "physical_meaning": (
            "Each per-device certificate artifact is tied to the frozen evidence bundle, "
            "QC decisions, coefficient state, traceability records, and report release status."
        ),
        "artifact_count": len(rows),
        "artifacts": rows,
    }
    manifest_path = output_dir / "per_device_certificate_manifest.json"
    manifest_path.write_text(json.dumps(manifest_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    hashes_path = output_dir / "per_device_certificate_artifact_hashes.csv"
    fieldnames = ("artifact_key", "device_id", "certificate_kind", "format", "path", "sha256", "size_bytes")
    with hashes_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)

    return {
        "per_device_certificate_manifest": manifest_path,
        "per_device_certificate_artifact_hashes": hashes_path,
    }


def _per_device_point_evidence_rows(
    open_flow_summary: Sequence[Mapping[str, Any]],
) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for row in open_flow_summary:
        device_id = _device_id_from_row(row) or str(row.get("analyzer_prefix") or "unknown")
        rows.append(
            {
                "analyzer_device_id": device_id,
                "analyzer_prefix": row.get("analyzer_prefix") or "",
                "component": str(row.get("component") or "").upper(),
                "calibratability_grade": row.get("point_calibratability_grade") or "",
                "fit_input_role": row.get("point_calibratability_role") or "",
                "time_optimization_action": row.get("time_optimization_action") or "",
                "candidate_fit_allowed": _open_flow_fit_allowed(row),
                "candidate_fit_blockers": row.get("candidate_fit_blockers") or "",
                "sample_readiness_status": row.get("sample_readiness_status") or "",
                "sample_readiness_blockers": row.get("sample_readiness_blockers") or "",
                "queue_exclusion_status": row.get("queue_exclusion_status") or "",
                "a_grade_count": row.get("a_grade_count") or 0,
                "b_grade_count": row.get("b_grade_count") or 0,
                "rejected_count": row.get("rejected_count") or 0,
                "physical_meaning": (
                    "A=direct fit; B=stable state-normalized/review input; "
                    "C=not a formal fit input from the current evidence."
                ),
            }
        )
    return sorted(rows, key=lambda item: (str(item["analyzer_device_id"]), str(item["component"])))


def _device_ids_from_evidence(*groups: Sequence[Mapping[str, Any]]) -> List[str]:
    ids: List[str] = []
    for rows in groups:
        for row in rows:
            device_id = _device_id_from_row(row)
            if (
                _is_single_report_analyzer_id(device_id)
                and device_id not in ids
            ):
                ids.append(device_id)
    return sorted(ids)


def _component_set(rows: Sequence[Mapping[str, Any]]) -> str:
    values = sorted(
        {
            str(row.get("component") or "").upper()
            for row in rows
            if str(row.get("component") or "").strip()
        }
    )
    return ";".join(values)


def _status_for_certificate(
    *,
    formal_issue_allowed: bool,
    point_rows: Sequence[Mapping[str, Any]],
    candidate_rows: Sequence[Mapping[str, Any]],
    release_status: str,
) -> tuple[str, List[str]]:
    reasons: List[str] = []
    if not point_rows:
        reasons.append("open_flow_point_evidence_missing")
    if any(str(row.get("calibratability_grade") or "").upper() == "C" for row in point_rows):
        reasons.append("contains_non_calibratable_points")
    if any(not _bool_value(row.get("candidate_fit_allowed")) for row in point_rows):
        reasons.append("open_flow_candidate_fit_blocked")
    if any(str(row.get("consolidated_status") or "") == "blocked" for row in candidate_rows):
        reasons.append("coefficient_candidate_blocked")
    if not formal_issue_allowed:
        reasons.append(f"report_release_status={release_status or 'not_available'}")
    if reasons:
        return "draft_or_blocked", reasons
    return "ready_for_formal_issue", ["all_device_certificate_gates_passed"]


def _status_for_verification_certificate(
    *,
    formal_issue_allowed: bool,
    post_write_rows: Sequence[Mapping[str, Any]],
    release_status: str,
) -> tuple[str, List[str]]:
    if not post_write_rows:
        return "not_available", ["post_write_or_independent_verification_missing"]
    statuses = {str(row.get("status") or "").strip().lower() for row in post_write_rows}
    reasons: List[str] = []
    if statuses - {"pass"}:
        reasons.append("verification_points_not_all_pass")
    if not formal_issue_allowed:
        reasons.append(f"report_release_status={release_status or 'not_available'}")
    if reasons:
        return "draft_or_blocked", reasons
    return "ready_for_formal_issue", ["all_device_verification_gates_passed"]


def _per_device_certificate_readiness_rows(
    *,
    point_evidence_rows: Sequence[Mapping[str, Any]],
    candidate_review_rollup: Sequence[Mapping[str, Any]],
    coefficient_snapshot_rows: Sequence[Mapping[str, Any]],
    coefficient_write_event_rows: Sequence[Mapping[str, Any]],
    post_write_reverification: Mapping[str, Any],
    release_decision: Mapping[str, Any],
) -> List[Dict[str, Any]]:
    post_write_rows = [
        dict(row) for row in post_write_reverification.get("device_summary") or []
    ]
    device_ids = _device_ids_from_evidence(
        point_evidence_rows,
        candidate_review_rollup,
        coefficient_snapshot_rows,
        coefficient_write_event_rows,
        post_write_rows,
    )
    formal_issue_allowed = bool(release_decision.get("formal_issue_allowed"))
    release_status = str(release_decision.get("release_status") or "")
    rows: List[Dict[str, Any]] = []
    for device_id in device_ids:
        points = [
            row for row in point_evidence_rows
            if str(row.get("analyzer_device_id") or "") == device_id
        ]
        candidates = [
            row for row in candidate_review_rollup
            if str(row.get("analyzer_device_id") or "") == device_id
        ]
        snapshots = [
            row for row in coefficient_snapshot_rows
            if str(row.get("analyzer_id") or "") == device_id
        ]
        write_events = [
            row for row in coefficient_write_event_rows
            if str(row.get("analyzer_id") or "") == device_id
        ]
        verification_rows = [
            row for row in post_write_rows
            if str(row.get("device_id") or row.get("analyzer_device_id") or "") == device_id
        ]
        calibration_status, calibration_reasons = _status_for_certificate(
            formal_issue_allowed=formal_issue_allowed,
            point_rows=points,
            candidate_rows=candidates,
            release_status=release_status,
        )
        verification_status, verification_reasons = _status_for_verification_certificate(
            formal_issue_allowed=formal_issue_allowed,
            post_write_rows=verification_rows,
            release_status=release_status,
        )
        rows.append(
            {
                "analyzer_device_id": device_id,
                "components": _component_set(points),
                "point_evidence_count": len(points),
                "calibratable_A_count": sum(
                    1 for row in points if str(row.get("calibratability_grade") or "").upper() == "A"
                ),
                "calibratable_B_count": sum(
                    1 for row in points if str(row.get("calibratability_grade") or "").upper() == "B"
                ),
                "calibratable_C_count": sum(
                    1 for row in points if str(row.get("calibratability_grade") or "").upper() == "C"
                ),
                "candidate_components": _component_set(candidates),
                "candidate_statuses": ";".join(
                    sorted({str(row.get("consolidated_status") or "") for row in candidates if row.get("consolidated_status")})
                ),
                "coefficient_snapshot_count": len(snapshots),
                "coefficient_write_event_count": len(write_events),
                "post_write_verification_point_count": sum(
                    int(_safe_float(row.get("point_count")) or 0) for row in verification_rows
                ),
                "post_write_verification_statuses": ";".join(
                    sorted({str(row.get("status") or "") for row in verification_rows if row.get("status")})
                ),
                "calibration_certificate_status": calibration_status,
                "calibration_certificate_reasons": ";".join(calibration_reasons),
                "verification_certificate_status": verification_status,
                "verification_certificate_reasons": ";".join(verification_reasons),
                "formal_issue_allowed": formal_issue_allowed,
                "release_status": release_status,
            }
        )
    return rows


def _summarize_points(calibration_points: Sequence[Mapping[str, Any]]) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for point in calibration_points:
        rows.append(
            {
                "component": str(point.get("component") or "").upper(),
                "point_id": point.get("point_key") or "",
                "pressure_mode": point.get("pressure_mode") or "",
                "samples": point.get("sample_count") or 0,
                "a_grade": point.get("a_grade_count") or 0,
                "b_grade": point.get("b_grade_count") or 0,
                "rejected": point.get("rejected_count") or 0,
            }
        )
    return rows


def _uncertainty_budget(
    *,
    standard_gases: Sequence[Mapping[str, Any]],
    results: Sequence[Mapping[str, Any]],
    pressure_qc: Sequence[Mapping[str, Any]],
) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for component in ("co2", "h2o"):
        cert_u = _component_uncertainty(standard_gases, component)
        component_results = [row for row in results if str(row.get("component") or "").lower() == component]
        stdevs = [_safe_float(row.get("std_dev")) for row in component_results]
        repeatability = max([value for value in stdevs if value is not None], default=None)
        rows.extend(
            [
                {
                    "component": component.upper(),
                    "source": "standard_gas_certificate",
                    "status": "evaluated" if cert_u is not None else "not_evaluated",
                    "value": _fmt(cert_u),
                    "unit": "certificate_unit",
                    "basis": "standard gas certificate snapshot",
                },
                {
                    "component": component.upper(),
                    "source": "repeatability",
                    "status": "estimated_from_a_grade_samples" if repeatability is not None else "not_evaluated",
                    "value": _fmt(repeatability),
                    "unit": "component_unit",
                    "basis": "A-grade sample standard deviation",
                },
                {
                    "component": component.upper(),
                    "source": "fit_residual",
                    "status": "not_evaluated",
                    "value": "not_released",
                    "unit": "component_unit",
                    "basis": "candidate coefficient solver qualification pending",
                },
                {
                    "component": component.upper(),
                    "source": "analyzer_resolution",
                    "status": "not_evaluated",
                    "value": "not_released",
                    "unit": "component_unit",
                    "basis": "resolution model not yet released",
                },
            ]
        )

    pressure_metrics = pressure_qc[0] if pressure_qc else {}
    pressure_bias = _safe_float(
        pressure_metrics.get("analyzer_minus_com22_max_abs_hpa")
        or pressure_metrics.get("analyzer_minus_com22_mean_hpa")
    )
    shared_rows = [
        {
            "component": "CO2/H2O",
            "source": "pressure_channel_bias",
            "status": "evaluated" if pressure_bias is not None else "not_evaluated",
            "value": _fmt(pressure_bias),
            "unit": "hPa",
            "basis": "analyzer pressure P versus COM22 quick check",
        },
        {
            "component": "H2O",
            "source": "dewpoint_or_humidity_reference",
            "status": "not_evaluated",
            "value": "not_released",
            "unit": "humidity_unit",
            "basis": "dewpoint/humidity reference certificate chain must be attached",
        },
        {
            "component": "CO2/H2O",
            "source": "temperature_effect",
            "status": "not_evaluated",
            "value": "not_released",
            "unit": "component_unit",
            "basis": "temperature sensitivity model pending",
        },
        {
            "component": "CO2/H2O",
            "source": "sampling_stability",
            "status": "evaluated",
            "value": "QC_gate",
            "unit": "qualitative",
            "basis": "open-flow stability/QC classification",
        },
        {
            "component": "H2O",
            "source": "water_vapor_correction",
            "status": "not_evaluated",
            "value": "not_released",
            "unit": "component_unit",
            "basis": "dry/wet conversion uncertainty model pending",
        },
    ]
    rows.extend(shared_rows)
    for row in rows:
        quantity = str(row.get("source") or "")
        value = _safe_float(row.get("value"))
        sensitivity = 1.0 if value is not None else None
        contribution = value * sensitivity if value is not None and sensitivity is not None else None
        row.setdefault("input_quantity", quantity)
        row.setdefault("distribution", "normal_or_certificate")
        row.setdefault("standard_uncertainty", value)
        row.setdefault("sensitivity_coefficient", sensitivity)
        row.setdefault("contribution", contribution)
        row.setdefault("evidence_source", row.get("basis") or "")
        row.setdefault("missing_reason", "" if contribution is not None else str(row.get("status") or "not_evaluated"))
    return rows


REQUIRED_UNCERTAINTY_BY_COMPONENT = {
    "CO2": {
        "standard_gas_certificate",
        "repeatability",
        "fit_residual",
        "analyzer_resolution",
        "pressure_channel_bias",
        "temperature_effect",
        "sampling_stability",
    },
    "H2O": {
        "standard_gas_certificate",
        "repeatability",
        "fit_residual",
        "analyzer_resolution",
        "pressure_channel_bias",
        "dewpoint_or_humidity_reference",
        "temperature_effect",
        "sampling_stability",
        "water_vapor_correction",
    },
}


def _load_optional_json(path: str | Path | None) -> Dict[str, Any]:
    if not path:
        return {}
    source = Path(path)
    if not source.exists():
        raise FileNotFoundError(f"JSON file not found: {source}")
    return json.loads(source.read_text(encoding="utf-8"))


def _normalize_uncertainty_input(row: Mapping[str, Any]) -> Dict[str, Any]:
    quantity = str(row.get("input_quantity") or row.get("source") or "")
    component = str(row.get("component") or "CO2/H2O").upper()
    standard_uncertainty = _safe_float(row.get("standard_uncertainty", row.get("value")))
    sensitivity = _safe_float(row.get("sensitivity_coefficient", 1.0))
    contribution = _safe_float(row.get("contribution"))
    if contribution is None and standard_uncertainty is not None and sensitivity is not None:
        contribution = standard_uncertainty * sensitivity
    status = str(row.get("status") or ("released" if contribution is not None else "not_evaluated"))
    return {
        "component": component,
        "source": quantity,
        "input_quantity": quantity,
        "distribution": str(row.get("distribution") or "normal"),
        "status": status,
        "value": _fmt(standard_uncertainty),
        "standard_uncertainty": standard_uncertainty,
        "sensitivity_coefficient": sensitivity,
        "contribution": contribution,
        "unit": str(row.get("unit") or "component_unit"),
        "basis": str(row.get("basis") or row.get("evidence_source") or ""),
        "evidence_source": str(row.get("evidence_source") or row.get("basis") or ""),
        "missing_reason": "" if contribution is not None or quantity == "sampling_stability" else "missing_numeric_contribution",
    }


def _merge_uncertainty_inputs(
    default_rows: Sequence[Mapping[str, Any]],
    uncertainty_payload: Mapping[str, Any],
) -> List[Dict[str, Any]]:
    merged = [dict(row) for row in default_rows]
    inputs = uncertainty_payload.get("inputs") if isinstance(uncertainty_payload, Mapping) else None
    if not isinstance(inputs, Sequence) or isinstance(inputs, (str, bytes)):
        return merged

    by_key: Dict[tuple[str, str], Dict[str, Any]] = {}
    order: List[tuple[str, str]] = []
    for row in merged:
        key = (str(row.get("component") or "CO2/H2O").upper(), str(row.get("input_quantity") or row.get("source") or ""))
        by_key[key] = row
        order.append(key)
    for source in inputs:
        if not isinstance(source, Mapping):
            continue
        row = _normalize_uncertainty_input(source)
        key = (str(row.get("component") or "CO2/H2O").upper(), str(row.get("input_quantity") or ""))
        if key not in by_key:
            order.append(key)
        by_key[key] = row
    return [by_key[key] for key in order]


def _uncertainty_row_for(
    rows: Sequence[Mapping[str, Any]],
    *,
    component: str,
    quantity: str,
) -> Optional[Mapping[str, Any]]:
    allowed_components = {component, "CO2/H2O"}
    for row in rows:
        row_component = str(row.get("component") or "").upper()
        row_quantity = str(row.get("input_quantity") or row.get("source") or "")
        if row_component in allowed_components and row_quantity == quantity:
            return row
    return None


def _evaluate_uncertainty_release(
    rows: Sequence[Mapping[str, Any]],
    uncertainty_payload: Mapping[str, Any],
) -> Dict[str, Any]:
    coverage_factor = _safe_float(uncertainty_payload.get("coverage_factor")) or 2.0
    explicit_release = bool(uncertainty_payload.get("released")) or str(
        uncertainty_payload.get("release_status") or ""
    ).lower() == "released"
    component_summaries: List[Dict[str, Any]] = []
    missing: List[str] = []

    for component, quantities in REQUIRED_UNCERTAINTY_BY_COMPONENT.items():
        contributions: List[float] = []
        for quantity in sorted(quantities):
            row = _uncertainty_row_for(rows, component=component, quantity=quantity)
            if row is None:
                missing.append(f"{component}:{quantity}:missing")
                continue
            contribution = _safe_float(row.get("contribution"))
            status = str(row.get("status") or "")
            if quantity == "sampling_stability" and contribution is None and status in {"released", "evaluated"}:
                contribution = 0.0
            if contribution is None:
                missing.append(f"{component}:{quantity}:{row.get('missing_reason') or status or 'not_evaluated'}")
                continue
            if status not in {"released", "evaluated", "estimated_from_a_grade_samples"}:
                missing.append(f"{component}:{quantity}:status={status}")
                continue
            contributions.append(float(contribution))
        combined = math.sqrt(sum(value * value for value in contributions)) if contributions else None
        component_summaries.append(
            {
                "component": component,
                "combined_standard_uncertainty": combined,
                "coverage_factor": coverage_factor,
                "expanded_uncertainty_k2": None if combined is None else combined * coverage_factor,
                "contribution_count": len(contributions),
            }
        )

    status = "released" if explicit_release and not missing else "not_released"
    return {
        "status": status,
        "released": status == "released",
        "coverage_factor": coverage_factor,
        "missing_required": missing,
        "component_summaries": component_summaries,
        "release_basis": str(uncertainty_payload.get("release_basis") or ""),
    }


def _apply_uncertainty_to_results(
    results: Sequence[Mapping[str, Any]],
    summary: Mapping[str, Any],
) -> List[Dict[str, Any]]:
    expanded_by_component = {
        str(row.get("component") or ""): row.get("expanded_uncertainty_k2")
        for row in summary.get("component_summaries", [])
        if isinstance(row, Mapping)
    }
    updated: List[Dict[str, Any]] = []
    for row in results:
        item = dict(row)
        value = _safe_float(expanded_by_component.get(str(item.get("component") or "")))
        item["expanded_uncertainty_k2"] = _fmt(value) if summary.get("released") and value is not None else "not_released"
        updated.append(item)
    return updated


def _decision(bundle: Mapping[str, Any]) -> Dict[str, Any]:
    run_rows = _table_rows(bundle, "runs")
    run = run_rows[0] if run_rows else {}
    candidates = _table_rows(bundle, "coefficient_candidates")
    write_events = _table_rows(bundle, "coefficient_write_events")
    evidence_status = str(run.get("evidence_status") or "")
    package_status = str(run.get("package_status") or "")
    if evidence_status == "blocked" or package_status == "blocked":
        status = "blocked"
    elif any(str(row.get("status") or "") not in {"not_attempted", "blocked"} for row in write_events):
        status = "coefficients_written_review_required"
    elif candidates:
        status = "candidate_coefficients_generated_no_write"
    else:
        status = "diagnostic_only"
    return {
        "decision_status": status,
        "evidence_status": evidence_status,
        "package_status": package_status,
        "package_blockers": run.get("package_blockers") or [],
    }


def _write_events_release_issue(write_events: Sequence[Mapping[str, Any]]) -> str:
    audited_statuses = {
        "not_attempted",
        "blocked",
        "written_readback_verified",
        "rollback_readback_verified",
    }
    for row in write_events:
        status = str(row.get("status") or "")
        if status not in audited_statuses:
            return f"coefficient_write_event_requires_audit:{status}"
    return ""


def _write_attempts_present(write_events: Sequence[Mapping[str, Any]]) -> bool:
    return any(str(row.get("status") or "") not in {"not_attempted", "blocked"} for row in write_events)


def _post_write_reverification_release_issue(
    write_events: Sequence[Mapping[str, Any]],
    post_write_reverification: Mapping[str, Any],
) -> str:
    if not _write_attempts_present(write_events):
        return ""
    if not post_write_reverification.get("available"):
        return "post_write_reverification_missing_after_coefficient_write"
    if str(post_write_reverification.get("overall_status") or "").lower() != "pass":
        return f"post_write_reverification_not_pass:{post_write_reverification.get('overall_status')}"
    return ""


def _pressure_release_issue(pressure_summary: Sequence[Mapping[str, Any]]) -> str:
    if not pressure_summary:
        return "pressure_channel_validation_missing"
    row = pressure_summary[0]
    if not _bool_value(row.get("allowed_for_co2_h2o_formal_work")):
        return "pressure_channel_not_formal_pass"
    return ""


def _signature_issue(reviewer: str, approver: str) -> str:
    missing: List[str] = []
    if not reviewer or reviewer == "pending_review":
        missing.append("reviewer_missing")
    if not approver or approver == "pending_approval":
        missing.append("approver_missing")
    return ";".join(missing)


def _release_checklist(
    *,
    decision: Mapping[str, Any],
    pressure_summary: Sequence[Mapping[str, Any]],
    uncertainty_summary: Mapping[str, Any],
    write_events: Sequence[Mapping[str, Any]],
    post_write_reverification: Mapping[str, Any],
    reviewer: str,
    approver: str,
) -> List[Dict[str, Any]]:
    write_issue = _write_events_release_issue(write_events)
    post_write_issue = _post_write_reverification_release_issue(write_events, post_write_reverification)
    pressure_issue = _pressure_release_issue(pressure_summary)
    signature_issue = _signature_issue(reviewer, approver)
    return [
        {
            "check": "evidence_package_not_blocked",
            "status": "pass" if decision.get("decision_status") != "blocked" else "fail",
            "reason": ";".join(_split_reasons(decision.get("package_blockers"))) or "",
        },
        {
            "check": "pressure_channel_formal_pass",
            "status": "pass" if not pressure_issue else "fail",
            "reason": pressure_issue,
        },
        {
            "check": "uncertainty_budget_released",
            "status": "pass" if uncertainty_summary.get("released") else "fail",
            "reason": ";".join(uncertainty_summary.get("missing_required") or []),
        },
        {
            "check": "coefficient_write_audited_or_no_write",
            "status": "pass" if not write_issue else "fail",
            "reason": write_issue,
        },
        {
            "check": "post_write_reverification_passed_or_no_write",
            "status": "pass" if not post_write_issue else "fail",
            "reason": post_write_issue,
        },
        {
            "check": "reviewer_and_approver_present",
            "status": "pass" if not signature_issue else "fail",
            "reason": signature_issue,
        },
    ]


def _report_release_decision(
    *,
    decision: Mapping[str, Any],
    pressure_summary: Sequence[Mapping[str, Any]],
    uncertainty_summary: Mapping[str, Any],
    write_events: Sequence[Mapping[str, Any]],
    post_write_reverification: Mapping[str, Any],
    reviewer: str,
    approver: str,
) -> Dict[str, Any]:
    if decision.get("decision_status") == "blocked":
        status = "blocked"
        reasons = _split_reasons(decision.get("package_blockers")) or ["evidence_package_blocked"]
    else:
        write_issue = _write_events_release_issue(write_events)
        post_write_issue = _post_write_reverification_release_issue(write_events, post_write_reverification)
        pressure_issue = _pressure_release_issue(pressure_summary)
        signature_issue = _signature_issue(reviewer, approver)
        reasons = []
        if write_issue:
            status = "not_releasable"
            reasons.append(write_issue)
        elif post_write_issue:
            status = "not_releasable"
            reasons.append(post_write_issue)
        elif pressure_issue:
            status = "blocked"
            reasons.append(pressure_issue)
        elif not uncertainty_summary.get("released"):
            status = "draft_only"
            reasons.append("uncertainty_budget_not_released")
        elif signature_issue:
            status = "review_ready"
            reasons.extend(_split_reasons(signature_issue))
        else:
            status = "formal_release_ready"
            reasons.append("all_release_gates_passed")
    return {
        "release_status": status,
        "issue_mark": "" if status == "formal_release_ready" else "DRAFT / NOT FOR FORMAL ISSUE",
        "reasons": reasons,
        "formal_issue_allowed": status == "formal_release_ready",
    }


def _build_report_model_from_bundle_base(
    bundle: Mapping[str, Any],
    *,
    report_no: str = "",
    reviewer: str = "",
    approver: str = "",
    location: str = "",
    calibration_date: str = "",
    analyzer_prefix: str = "ga01",
    uncertainty_payload: Optional[Mapping[str, Any]] = None,
) -> Dict[str, Any]:
    run_rows = _table_rows(bundle, "runs")
    run = run_rows[0] if run_rows else {}
    standard_gases = _table_rows(bundle, "standard_gases")
    reference_certificates = _table_rows(bundle, "reference_certificates")
    calibration_points = _table_rows(bundle, "calibration_points")
    sample_files = _table_rows(bundle, "sample_files")
    qc_results = _table_rows(bundle, "qc_results")
    coefficient_snapshots = _table_rows(bundle, "coefficient_snapshots")
    candidates = _table_rows(bundle, "coefficient_candidates")
    write_events = _table_rows(bundle, "coefficient_write_events")
    integrity_checks = _table_rows(bundle, "evidence_integrity_checks")
    a_grade_rows = _artifact_rows(bundle, "a_grade_samples.csv")
    rejected_rows = _artifact_rows(bundle, "rejected_samples.csv")
    pressure_summary = _artifact_rows(bundle, "pressure_validation_summary.csv")
    open_flow_summary = _artifact_rows(bundle, "open_flow_run_summary.csv")
    h2o_queue_exclusions = _h2o_queue_exclusion_rows(bundle)
    open_flow_summary = _apply_h2o_queue_exclusions(open_flow_summary, h2o_queue_exclusions)
    post_write_reverification = _post_write_reverification_model(bundle)
    run_evidence_status = _run_evidence_status_model(bundle)
    results = _group_component_results(a_grade_rows, standard_gases, analyzer_prefix=analyzer_prefix)
    decision = _decision(bundle)
    candidate_review_rollup = _candidate_review_rollup_rows(candidates)
    coefficient_snapshot_rows = _coefficient_snapshot_rows(coefficient_snapshots)
    coefficient_write_event_rows = _coefficient_write_event_rows(write_events)
    base_uncertainty_budget = _uncertainty_budget(
        standard_gases=standard_gases,
        results=results,
        pressure_qc=pressure_summary,
    )
    uncertainty_inputs = dict(uncertainty_payload or {})
    uncertainty_budget = _merge_uncertainty_inputs(base_uncertainty_budget, uncertainty_inputs)
    uncertainty_summary = _evaluate_uncertainty_release(uncertainty_budget, uncertainty_inputs)
    results = _apply_uncertainty_to_results(results, uncertainty_summary)
    normalized_reviewer = reviewer or "pending_review"
    normalized_approver = approver or "pending_approval"
    report_release_decision = _report_release_decision(
        decision=decision,
        pressure_summary=pressure_summary,
        uncertainty_summary=uncertainty_summary,
        write_events=write_events,
        post_write_reverification=post_write_reverification,
        reviewer=normalized_reviewer,
        approver=normalized_approver,
    )
    release_checklist = _release_checklist(
        decision=decision,
        pressure_summary=pressure_summary,
        uncertainty_summary=uncertainty_summary,
        write_events=write_events,
        post_write_reverification=post_write_reverification,
        reviewer=normalized_reviewer,
        approver=normalized_approver,
    )
    per_device_point_evidence = _per_device_point_evidence_rows(open_flow_summary)
    per_device_certificate_readiness = _per_device_certificate_readiness_rows(
        point_evidence_rows=per_device_point_evidence,
        candidate_review_rollup=candidate_review_rollup,
        coefficient_snapshot_rows=coefficient_snapshot_rows,
        coefficient_write_event_rows=coefficient_write_event_rows,
        post_write_reverification=post_write_reverification,
        release_decision=report_release_decision,
    )
    model = {
        "report_no": report_no or f"V15-{str(bundle.get('run_id') or run.get('run_id') or 'run')}",
        "run_id": str(bundle.get("run_id") or run.get("run_id") or ""),
        "run_db_id": str(bundle.get("run_db_id") or run.get("id") or ""),
        "generated_at": _now(),
        "calibration_date": calibration_date or str(run.get("created_at") or "")[:10] or datetime.now().date().isoformat(),
        "location": location or "not_specified",
        "analyzer_id": str(run.get("analyzer_id") or ""),
        "operator": str(run.get("operator_name") or ""),
        "reviewer": normalized_reviewer,
        "approver": normalized_approver,
        "decision": decision,
        "report_release_decision": report_release_decision,
        "release_checklist": release_checklist,
        "standard_gases": standard_gases,
        "reference_certificates": reference_certificates,
        "calibration_points": _summarize_points(calibration_points),
        "result_rows": results,
        "pressure_summary": pressure_summary,
        "open_flow_summary": open_flow_summary,
        "h2o_queue_exclusions": h2o_queue_exclusions,
        "post_write_reverification": post_write_reverification,
        "run_evidence_status": run_evidence_status,
        "qc_results": qc_results,
        "coefficient_snapshots": coefficient_snapshots,
        "coefficient_snapshot_rows": coefficient_snapshot_rows,
        "candidate_rows": candidates,
        "candidate_review_rollup": candidate_review_rollup,
        "h2o_candidate_review_rollup": [
            row for row in candidate_review_rollup if str(row.get("component") or "").lower() == "h2o"
        ],
        "write_events": write_events,
        "coefficient_write_event_rows": coefficient_write_event_rows,
        "per_device_point_evidence": per_device_point_evidence,
        "per_device_certificate_readiness": per_device_certificate_readiness,
        "integrity_checks": integrity_checks,
        "rejected_rows": rejected_rows[:200],
        "artifact_manifest": [
            {
                "role": row.get("artifact_role"),
                "path": row.get("path"),
                "sha256": row.get("sha256"),
                "required": row.get("required"),
            }
            for row in sample_files
        ],
    }
    model["uncertainty_budget"] = uncertainty_budget
    model["uncertainty_summary"] = uncertainty_summary
    model["scope_statement"] = (
        "本次 CO2/H2O 主校准基于开放流通、当前大气压附近条件。"
        "封路控压多压力点、PACE OUTPUT 长期开路动态控压、PACE ACT + sink bias、"
        "VENT-hold 均作为工程诊断边界处理，默认不进入正式 CO2/H2O 拟合。"
    )
    model["method_statement"] = (
        "标准气持续开放流通进入分析仪，持续刷新光学腔体和下游管路，"
        "旧气、死体积湿气和残余气被带走。稳定性门禁确认 CO2/H2O、压力、"
        "露点、温度和工厂模式 ratio/signal 在采样窗口内可用于复核。"
    )
    model["coefficient_statement"] = _coefficient_statement(write_events, decision)
    model["pressure_compensation_statement"] = _pressure_compensation_statement(pressure_summary)
    model["limitations"] = _limitations(decision, pressure_summary)
    return model


def _simple_rows(rows: Sequence[Mapping[str, Any]], keys: Sequence[str], limit: int = 20) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for row in list(rows)[:limit]:
        out.append({key: row.get(key, "") for key in keys})
    return out


def _metadata(row: Mapping[str, Any]) -> Mapping[str, Any]:
    value = row.get("metadata")
    return value if isinstance(value, Mapping) else {}


def _json_cell(value: Any) -> str:
    if value in (None, "", [], {}):
        return ""
    if isinstance(value, (Mapping, list, tuple)):
        return json.dumps(value, ensure_ascii=False, separators=(",", ":"), default=str)
    return str(value)


def _coefficient_payload(row: Mapping[str, Any]) -> Mapping[str, Any]:
    value = row.get("coefficients")
    return value if isinstance(value, Mapping) else {}


def _getco_snapshot_value(coefficients: Mapping[str, Any], group: str) -> str:
    for key in (f"{group}_after", f"{group}_before", group):
        if key in coefficients:
            return _json_cell(coefficients.get(key))
    return ""


def _coefficient_snapshot_rows(snapshots: Sequence[Mapping[str, Any]]) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for row in snapshots:
        coefficients = _coefficient_payload(row)
        metadata = _metadata(row)
        getco_groups = metadata.get("getco_groups")
        if isinstance(getco_groups, (list, tuple)):
            getco_group_text = ";".join(str(item) for item in getco_groups)
        else:
            getco_group_text = str(getco_groups or "")
        snapshot_row: Dict[str, Any] = {
            "analyzer_id": row.get("analyzer_id") or _candidate_device_id(row),
            "snapshot_type": row.get("snapshot_type") or "",
            "getco_groups": getco_group_text,
            "coefficients_hash": row.get("coefficients_hash") or "",
            "snapshot_file": metadata.get("snapshot_file_name") or Path(str(metadata.get("path") or "")).name,
        }
        for index in range(1, 10):
            group = f"GETCO{index}"
            snapshot_row[group] = _getco_snapshot_value(coefficients, group)
        rows.append(snapshot_row)
    return rows


def _coefficient_write_event_rows(write_events: Sequence[Mapping[str, Any]]) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for row in write_events:
        rows.append(
            {
                "analyzer_id": row.get("analyzer_id") or "",
                "event_type": row.get("event_type") or "",
                "status": row.get("status") or "",
                "candidate_id": row.get("candidate_id") or "",
                "old_coefficients_hash": row.get("old_coefficients_hash") or "",
                "approved_by": row.get("approved_by") or "",
                "command_summary": row.get("command_summary") or "",
                "readback": _json_cell(row.get("readback")),
            }
        )
    return rows


def _candidate_device_id(row: Mapping[str, Any]) -> str:
    for key in ("analyzer_device_id", "device_id", "analyzer_id", "sensor_id"):
        value = row.get(key)
        if value not in (None, ""):
            return str(value)
    metadata = _metadata(row)
    for key in ("analyzer_device_id", "device_id", "analyzer_id", "sensor_id"):
        value = metadata.get(key)
        if value not in (None, ""):
            return str(value)
    return "unknown"


def _reason_list(value: Any) -> List[str]:
    if value in (None, ""):
        return []
    if isinstance(value, list):
        return [str(item) for item in value if str(item)]
    if isinstance(value, tuple):
        return [str(item) for item in value if str(item)]
    return [part for part in str(value).split(";") if part]


def _candidate_warning_reasons(row: Mapping[str, Any]) -> List[str]:
    warnings = _reason_list(row.get("warnings"))
    metadata = _metadata(row)
    warnings.extend(_reason_list(metadata.get("warning_reasons_list")))
    warnings.extend(_reason_list(metadata.get("warning_reasons")))
    return sorted(set(warnings))


def _candidate_blockers(row: Mapping[str, Any]) -> List[str]:
    blockers = _reason_list(row.get("blockers"))
    metadata = _metadata(row)
    blockers.extend(_reason_list(metadata.get("blocked_reasons")))
    return sorted(set(blockers))


def _candidate_rollup_status(statuses: Sequence[str], blockers: Sequence[str]) -> str:
    lowered = {str(status or "").strip().lower() for status in statuses if str(status or "").strip()}
    if blockers or "blocked" in lowered:
        return "blocked"
    if any("review" in status for status in lowered):
        return "review_required"
    if lowered <= {"ready_for_reviewer", "ready", "pass"}:
        return "ready_for_reviewer"
    return "review_required" if lowered else "not_available"


def _candidate_rollup_action(status: str) -> str:
    if status == "blocked":
        return "先处理阻断原因，不得写入"
    if status == "review_required":
        return "允许人工评审，不得自动写入"
    if status == "ready_for_reviewer":
        return "可进入审核员评审，不得自动写入"
    return "补充候选证据"


def _candidate_review_rollup_rows(candidates: Sequence[Mapping[str, Any]]) -> List[Dict[str, Any]]:
    grouped: Dict[tuple[str, str], Dict[str, Any]] = {}
    for row in candidates:
        component = str(row.get("component") or "unknown").strip().lower() or "unknown"
        device_id = _candidate_device_id(row)
        key = (component, device_id)
        item = grouped.setdefault(
            key,
            {
                "component": component,
                "analyzer_device_id": device_id,
                "source_row_count": 0,
                "statuses": [],
                "blockers": [],
                "warnings": [],
            },
        )
        item["source_row_count"] += 1
        item["statuses"].append(str(row.get("candidate_status") or ""))
        item["blockers"].extend(_candidate_blockers(row))
        item["warnings"].extend(_candidate_warning_reasons(row))

    rows: List[Dict[str, Any]] = []
    for item in grouped.values():
        statuses = sorted(set(item["statuses"]))
        blockers = sorted(set(item["blockers"]))
        warnings = sorted(set(item["warnings"]))
        status = _candidate_rollup_status(statuses, blockers)
        rows.append(
            {
                "component": item["component"],
                "analyzer_device_id": item["analyzer_device_id"],
                "source_row_count": item["source_row_count"],
                "statuses": ";".join(statuses),
                "consolidated_status": status,
                "blockers": ";".join(blockers),
                "warnings": ";".join(warnings),
                "review_action": _candidate_rollup_action(status),
            }
        )
    return sorted(rows, key=lambda row: (str(row["component"]), str(row["analyzer_device_id"])))


def _bool_label(value: Any) -> str:
    return "是" if _bool_value(value) else "否"


def _run_evidence_status_paragraphs(status: Mapping[str, Any]) -> List[str]:
    if not status.get("available"):
        return [
            "运行证据状态文件未进入当前证据包；报告只能依据 evidence_bundle 中已索引的表格给出有限结论。"
        ]
    boundaries = status.get("physical_boundaries")
    if not isinstance(boundaries, Mapping):
        boundaries = {}
    source_names = [
        Path(str(path)).name
        for path in status.get("source_artifacts") or []
        if str(path)
    ]
    source_text = "；".join(source_names) if source_names else "not_available"
    return [
        f"运行证据总状态：{status.get('overall_status') or 'not_available'}",
        f"当前阶段：{status.get('current_stage') or 'not_available'}",
        f"流程合同状态：{status.get('contract_status') or 'not_available'}",
        f"来源状态文件：{source_text}",
        (
            "物理边界："
            f"仅离线状态={_bool_label(boundaries.get('offline_status_only'))}；"
            f"打开 COM={_bool_label(boundaries.get('opens_com_ports'))}；"
            f"控制气路/水路={_bool_label(boundaries.get('controls_water_or_gas_routes'))}；"
            f"控制阀或 PACE={_bool_label(boundaries.get('controls_valves_or_pace'))}；"
            f"写入系数={_bool_label(boundaries.get('writes_coefficients'))}；"
            f"非 real acceptance 证据={_bool_label(boundaries.get('not_real_acceptance_evidence'))}。"
        ),
    ]


def _run_evidence_stage_rows(status: Mapping[str, Any], limit: int = 20) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for stage in list(status.get("stage_statuses") or [])[:limit]:
        if not isinstance(stage, Mapping):
            continue
        artifact_roles = stage.get("artifact_roles") or ()
        if isinstance(artifact_roles, str):
            artifact_roles_text = artifact_roles
        else:
            artifact_roles_text = ";".join(str(role) for role in artifact_roles if str(role))
        rows.append(
            {
                "阶段ID": stage.get("stage_id", ""),
                "阶段名称": stage.get("title", ""),
                "状态": stage.get("status", ""),
                "原因": stage.get("reason", ""),
                "证据角色": artifact_roles_text,
                "证据数量": stage.get("artifact_count", ""),
                "物理意义": stage.get("physical_meaning", ""),
            }
        )
    return rows


def _run_evidence_boundary_rows(status: Mapping[str, Any]) -> List[Dict[str, Any]]:
    boundaries = status.get("physical_boundaries")
    if not isinstance(boundaries, Mapping):
        boundaries = {}
    labels = [
        ("offline_status_only", "仅离线状态"),
        ("opens_com_ports", "打开 COM"),
        ("controls_water_or_gas_routes", "控制气路/水路"),
        ("controls_valves_or_pace", "控制阀或 PACE"),
        ("writes_coefficients", "写入系数"),
        ("not_real_acceptance_evidence", "非 real acceptance 证据"),
    ]
    return [{"边界": label, "值": _bool_label(boundaries.get(key))} for key, label in labels]


def _scope_statement() -> str:
    return (
        "本次 CO2/H2O 主校准基于开放流通、当前大气压附近条件。"
        "封路控压多压力点、PACE OUTPUT 长期开路动态控压、PACE ACT + sink bias、"
        "VENT-hold 均作为工程诊断边界处理，默认不进入正式 CO2/H2O 拟合。"
    )


def _method_statement() -> str:
    return (
        "标准气持续开放流通进入分析仪，持续刷新光学腔体和下游管路，"
        "旧气、死体积湿气和残余气被带走。稳定性门禁确认 CO2/H2O、压力、"
        "露点、温度和工厂模式 ratio/signal 在采样窗口内可用于复核。"
    )


def _coefficient_statement(write_events: Sequence[Mapping[str, Any]], decision: Mapping[str, Any]) -> str:
    if not write_events:
        return "本次未发现系数写入事件。候选系数如有生成，也只能进入评审，不能自动写入设备。"
    statuses = {str(row.get("status") or "") for row in write_events}
    if statuses <= {"not_attempted", "blocked"}:
        return (
            "本次仅生成候选系数或候选评审材料，未写入设备。"
            "SENCOx 写入仍需要单独授权、旧系数备份、写入后读回和独立复验。"
        )
    return (
        "本次存在系数写入相关事件。报告必须同时复核旧系数、写入命令、读回值、"
        "审批记录、回滚证据以及写后独立复验结果。"
    )


def _pressure_compensation_statement(pressure_summary: Sequence[Mapping[str, Any]]) -> str:
    if not pressure_summary:
        return "压力通道快速验证证据缺失；压力补偿验证不在本次正式覆盖范围内。"
    row = pressure_summary[0]
    if _bool_value(row.get("allowed_for_co2_h2o_formal_work")):
        return (
            "分析仪内部压力 P 已通过当前大气压快速验证；"
            "多压力压力补偿验证仍为后置、可选的独立验证范围。"
        )
    return "分析仪内部压力 P 未达到正式 CO2/H2O 工作放行条件；压力补偿验证不得解释为正式覆盖。"


def _limitations(decision: Mapping[str, Any], pressure_summary: Sequence[Mapping[str, Any]]) -> List[str]:
    out = [
        "封路控压多压力点未用于正式 CO2/H2O 拟合。",
        "诊断数据可以保留，但不得作为 real acceptance 或正式拟合输入。",
        "不确定度预算第一版保留缺项状态；未释放的分量不得用于最终合格判定。",
    ]
    if decision.get("decision_status") == "blocked":
        out.append("证据包当前为 blocked，报告不得作为正式校准证书签发。")
    if not pressure_summary:
        out.append("压力通道快速验证缺失，压力补偿覆盖范围未建立。")
    return out


def build_report_model_from_bundle(
    bundle: Mapping[str, Any],
    *,
    report_no: str = "",
    reviewer: str = "",
    approver: str = "",
    location: str = "",
    calibration_date: str = "",
    analyzer_prefix: str = "ga01",
    uncertainty_payload: Optional[Mapping[str, Any]] = None,
) -> Dict[str, Any]:
    model = _build_report_model_from_bundle_base(
        bundle,
        report_no=report_no,
        reviewer=reviewer,
        approver=approver,
        location=location,
        calibration_date=calibration_date,
        analyzer_prefix=analyzer_prefix,
        uncertainty_payload=uncertainty_payload,
    )
    model["scope_statement"] = _scope_statement()
    model["method_statement"] = _method_statement()
    model["coefficient_statement"] = _coefficient_statement(
        model.get("write_events") or [],
        model.get("decision") or {},
    )
    model["pressure_compensation_statement"] = _pressure_compensation_statement(
        model.get("pressure_summary") or []
    )
    model["limitations"] = _limitations(
        model.get("decision") or {},
        model.get("pressure_summary") or [],
    )
    return model


def build_run_report(model: Mapping[str, Any]) -> ReportDocument:
    decision = model.get("decision") or {}
    release = model.get("report_release_decision") or {}
    post_write = model.get("post_write_reverification") or {}
    run_evidence_status = model.get("run_evidence_status") or {}
    return ReportDocument(
        title="V1.5 气体分析仪校准运行报告",
        sections=[
            ReportSection(
                "运行结论",
                [
                    f"报告编号：{model.get('report_no')}",
                    f"运行编号：{model.get('run_id')}",
                    f"证据状态：{decision.get('evidence_status')} / {decision.get('decision_status')}",
                    f"发布判定：{release.get('release_status')}",
                    str(release.get("issue_mark") or ""),
                    f"被校分析仪：{model.get('analyzer_id')}",
                    "本报告由证据包自动生成；报告生成过程不打开串口、不控制气路/水路、不写入设备。",
                ],
            ),
            ReportSection(
                "运行证据状态",
                _run_evidence_status_paragraphs(run_evidence_status),
                [
                    ReportTable(
                        "阶段状态",
                        _run_evidence_stage_rows(run_evidence_status, limit=12),
                    )
                ],
            ),
            ReportSection(
                "发布门禁",
                [
                    f"允许正式签发：{release.get('formal_issue_allowed')}",
                    f"阻断/提示原因：{';'.join(release.get('reasons') or [])}",
                ],
                [ReportTable("发布检查表", list(model.get("release_checklist") or []))],
            ),
            ReportSection(
                "写后复验",
                [
                    f"写后复验证据可用：{post_write.get('available')}",
                    f"写后复验状态：{post_write.get('overall_status')}",
                    "SENCO 写入成功和写后测量一致性是两个独立放行门禁。",
                ],
                [
                    ReportTable(
                        "写后复验设备汇总",
                        _simple_rows(
                            post_write.get("device_summary") or [],
                            (
                                "device_id",
                                "component",
                                "point_count",
                                "pass_count",
                                "fail_count",
                                "max_abs_error",
                                "max_abs_error_pct",
                                "status",
                            ),
                        ),
                    )
                ],
            ),
            ReportSection(
                "点位摘要",
                tables=[ReportTable("点位 QC 摘要", list(model.get("calibration_points") or []))],
            ),
            ReportSection(
                "每台设备证书准备状态",
                [
                    "本表按分析仪自身设备 ID 汇总，不以串口号作为身份。"
                    "校准证书和检定/复验证书只有在点位证据、候选系数、写入/复验、"
                    "不确定度和审核签名满足要求后才允许正式签发。"
                ],
                [
                    ReportTable(
                        "每台设备证书准备摘要",
                        list(model.get("per_device_certificate_readiness") or []),
                    )
                ],
            ),
            ReportSection(
                "候选系数评审",
                [model.get("coefficient_statement", "")],
                [
                    ReportTable(
                        "候选系数按设备归并",
                        _simple_rows(
                            model.get("candidate_review_rollup") or [],
                            (
                                "component",
                                "analyzer_device_id",
                                "source_row_count",
                                "consolidated_status",
                                "blockers",
                                "warnings",
                                "review_action",
                            ),
                            limit=80,
                        ),
                    ),
                    ReportTable(
                        "候选系数状态",
                        _simple_rows(
                            model.get("candidate_rows") or [],
                            (
                                "component",
                                "candidate_status",
                                "allowed_for_review",
                                "auto_write_allowed",
                                "blockers",
                            ),
                        ),
                    )
                ],
            ),
            ReportSection(
                "H2O 候选系数水路证据归并",
                [
                    "H2O 候选系数按设备独立评审；warning 只提示证据注意事项，不等同于 blocker。"
                    "水路候选不得被 CO2 成对写入门禁误判，也不得自动写入设备。"
                ],
                [
                    ReportTable(
                        "H2O 候选设备归并",
                        _simple_rows(
                            model.get("h2o_candidate_review_rollup") or [],
                            (
                                "analyzer_device_id",
                                "source_row_count",
                                "statuses",
                                "consolidated_status",
                                "blockers",
                                "warnings",
                                "review_action",
                            ),
                            limit=80,
                        ),
                    )
                ],
            ),
            ReportSection(
                "失败和限制",
                list(model.get("limitations") or []),
                [
                    ReportTable(
                        "完整性检查",
                        _simple_rows(
                            model.get("integrity_checks") or [],
                            ("check_name", "status", "severity", "details"),
                        ),
                    )
                ],
            ),
        ],
    )


def build_technical_report(model: Mapping[str, Any]) -> ReportDocument:
    release = model.get("report_release_decision") or {}
    post_write = model.get("post_write_reverification") or {}
    run_evidence_status = model.get("run_evidence_status") or {}
    return ReportDocument(
        title="V1.5 气体分析仪校准技术报告",
        sections=[
            ReportSection(
                "发布门禁",
                [
                    f"发布判定：{release.get('release_status')}",
                    str(release.get("issue_mark") or ""),
                    f"允许正式签发：{release.get('formal_issue_allowed')}",
                ],
                [ReportTable("发布检查表", list(model.get("release_checklist") or []))],
            ),
            ReportSection("方法和物理过程", [model.get("method_statement", "")]),
            ReportSection(
                "运行证据状态",
                _run_evidence_status_paragraphs(run_evidence_status),
                [
                    ReportTable("物理边界", _run_evidence_boundary_rows(run_evidence_status)),
                    ReportTable(
                        "阶段状态",
                        _run_evidence_stage_rows(run_evidence_status, limit=30),
                    ),
                ],
            ),
            ReportSection(
                "开放流通数据质量",
                [
                    "CO2/H2O 主校准只使用开放流通且组分稳定的数据。封路污染压力点、动态控压探针和 VENT-hold 证据保留为诊断边界。",
                    "工厂模式 ratio/signal、压力、温度和露点数据用于复核底层物理状态，不等同于全部进入正式拟合。",
                ],
                [
                    ReportTable(
                        "开放流通摘要",
                        _simple_rows(
                            model.get("open_flow_summary") or [],
                            (
                                "component",
                                "plan_status",
                                "pressure_channel_quick_check_status",
                                "a_grade_count",
                                "b_grade_count",
                                "rejected_count",
                                "candidate_fit_allowed",
                                "candidate_fit_blockers",
                            ),
                        ),
                    )
                ],
            ),
            ReportSection(
                "每台设备点位可校准性",
                [
                    "A 级点可直接作为拟合证据；B 级点表示状态稳定但需要归一化、桥接或人工评审；"
                    "C 级点不得进入正式拟合。该分级用于把气路/水路物理状态和证书签发门禁连接起来。"
                ],
                [
                    ReportTable(
                        "每台设备点位证据",
                        _simple_rows(
                            model.get("per_device_point_evidence") or [],
                            (
                                "analyzer_device_id",
                                "component",
                                "calibratability_grade",
                                "fit_input_role",
                                "time_optimization_action",
                                "candidate_fit_allowed",
                                "candidate_fit_blockers",
                                "sample_readiness_status",
                                "sample_readiness_blockers",
                            ),
                            limit=240,
                        ),
                    )
                ],
            ),
            ReportSection(
                "压力通道验证",
                [model.get("pressure_compensation_statement", "")],
                [
                    ReportTable(
                        "压力通道摘要",
                        _simple_rows(
                            model.get("pressure_summary") or [],
                            (
                                "status",
                                "validation_level",
                                "valid_pair_count",
                                "analyzer_minus_com22_mean_hpa",
                                "analyzer_minus_com22_max_abs_hpa",
                                "allowed_for_co2_h2o_formal_work",
                                "reason",
                            ),
                        ),
                    )
                ],
            ),
            ReportSection(
                "候选系数技术评审",
                [
                    "候选系数表保留原始候选行；按设备归并表用于审核员快速判断每台分析仪是否可评审、是否被阻断。"
                    "H2O warning 与 blocker 分离，避免把水路可解释提示当成正式阻断。"
                    "水路候选不得被 CO2 成对写入门禁误判，也不得自动写入设备。"
                ],
                [
                    ReportTable(
                        "候选系数按设备归并",
                        _simple_rows(
                            model.get("candidate_review_rollup") or [],
                            (
                                "component",
                                "analyzer_device_id",
                                "source_row_count",
                                "statuses",
                                "consolidated_status",
                                "blockers",
                                "warnings",
                                "review_action",
                            ),
                            limit=120,
                        ),
                    ),
                    ReportTable(
                        "H2O 候选设备归并",
                        _simple_rows(
                            model.get("h2o_candidate_review_rollup") or [],
                            (
                                "analyzer_device_id",
                                "source_row_count",
                                "statuses",
                                "consolidated_status",
                                "blockers",
                                "warnings",
                                "review_action",
                            ),
                            limit=120,
                        ),
                    ),
                ],
            ),
            ReportSection(
                "系数证据链",
                [
                    "GETCO1-9 快照冻结写入前和读回后的设备校准状态；它用于证明候选系数、受控写入、回滚和写后复验之间可以互相追溯。"
                    "SENCO 写入事件表只记录证据，不触发任何设备操作。"
                ],
                [
                    ReportTable(
                        "GETCO1-9 快照",
                        _simple_rows(
                            model.get("coefficient_snapshot_rows") or [],
                            (
                                "analyzer_id",
                                "snapshot_type",
                                "getco_groups",
                                "GETCO1",
                                "GETCO2",
                                "GETCO3",
                                "GETCO4",
                                "GETCO5",
                                "GETCO6",
                                "GETCO7",
                                "GETCO8",
                                "GETCO9",
                                "coefficients_hash",
                                "snapshot_file",
                            ),
                            limit=120,
                        ),
                    ),
                    ReportTable(
                        "系数写入事件",
                        _simple_rows(
                            model.get("coefficient_write_event_rows") or [],
                            (
                                "analyzer_id",
                                "event_type",
                                "status",
                                "candidate_id",
                                "old_coefficients_hash",
                                "approved_by",
                                "command_summary",
                                "readback",
                            ),
                            limit=120,
                        ),
                    ),
                ],
            ),
            ReportSection(
                "写后复验",
                [
                    f"写后复验证据可用：{post_write.get('available')}",
                    f"写后复验状态：{post_write.get('overall_status')}",
                    f"证据文件：{';'.join(post_write.get('source_artifacts') or [])}",
                ],
                [
                    ReportTable(
                        "写后复验点位明细",
                        _simple_rows(
                            post_write.get("point_results") or [],
                            (
                                "device_id",
                                "component",
                                "point_id",
                                "standard_value",
                                "measured_value",
                                "error",
                                "error_pct",
                                "limit_value",
                                "status",
                                "reason",
                            ),
                            limit=200,
                        ),
                    )
                ],
            ),
            ReportSection(
                "拒绝样本",
                ["拒绝样本必须保留原因，不能因为结果不好看而静默删除。"],
                [
                    ReportTable(
                        "拒绝样本前 200 行",
                        _simple_rows(
                            model.get("rejected_rows") or [],
                            ("component", "sample_index", "formal_grade", "formal_reject_reasons", "pressure_mode"),
                            limit=200,
                        ),
                    )
                ],
            ),
            ReportSection(
                "Artifact Hash",
                ["以下 hash 用于从证据包重建报告。"],
                [
                    ReportTable(
                        "证据文件",
                        _simple_rows(
                            model.get("artifact_manifest") or [],
                            ("role", "path", "sha256", "required"),
                            limit=200,
                        ),
                    )
                ],
            ),
        ],
    )


def build_formal_calibration_report(model: Mapping[str, Any]) -> ReportDocument:
    decision = model.get("decision") or {}
    release = model.get("report_release_decision") or {}
    post_write = model.get("post_write_reverification") or {}
    run_evidence_status = model.get("run_evidence_status") or {}
    cover = [
        f"报告发布判定：{release.get('release_status')}",
        str(release.get("issue_mark") or ""),
        f"写后复验可用：{post_write.get('available')}",
        f"写后复验状态：{post_write.get('overall_status')}",
        f"报告编号：{model.get('report_no')}",
        f"运行编号：{model.get('run_id')}",
        f"被校仪器：{model.get('analyzer_id')}",
        f"校准日期：{model.get('calibration_date')}",
        f"校准地点：{model.get('location')}",
        f"操作员：{model.get('operator')}",
        f"审核员：{model.get('reviewer')}",
        f"批准人：{model.get('approver')}",
        f"结论状态：{decision.get('decision_status')}",
    ]
    return ReportDocument(
        title="V1.5 气体分析仪正式校准报告",
        sections=[
            ReportSection("封面", cover),
            ReportSection(
                "发布门禁",
                [
                    f"允许正式签发：{release.get('formal_issue_allowed')}",
                    f"阻断/提示原因：{';'.join(release.get('reasons') or [])}",
                ],
                [ReportTable("发布检查表", list(model.get("release_checklist") or []))],
            ),
            ReportSection("校准范围声明", [model.get("scope_statement", ""), model.get("pressure_compensation_statement", "")]),
            ReportSection(
                "运行证据状态",
                _run_evidence_status_paragraphs(run_evidence_status),
                [
                    ReportTable(
                        "关键阶段状态",
                        _run_evidence_stage_rows(run_evidence_status, limit=20),
                    )
                ],
            ),
            ReportSection("校准方法", [model.get("method_statement", "")]),
            ReportSection(
                "标准和参考设备",
                [
                    "溯源围绕测量结果建立。标准气、压力、露点和温度参考应通过证书链和不确定度支撑结果有效性。",
                ],
                [
                    ReportTable(
                        "标准气",
                        _simple_rows(
                            model.get("standard_gases") or [],
                            (
                                "component",
                                "cylinder_id",
                                "certificate_value",
                                "certificate_uncertainty",
                                "valid_until",
                                "supplier",
                                "certificate_hash",
                            ),
                        ),
                    ),
                    ReportTable(
                        "参考设备证书",
                        _simple_rows(
                            model.get("reference_certificates") or [],
                            ("reference_role", "certificate_id", "certificate_hash", "valid_until", "uncertainty", "unit"),
                        ),
                    ),
                ],
            ),
            ReportSection(
                "数据质量摘要",
                tables=[ReportTable("点位质量", list(model.get("calibration_points") or []))],
            ),
            ReportSection(
                "每台设备证书状态",
                [
                    "正式校准证书和检定/复验证书必须按设备 ID 单独判断。"
                    "串口号只说明接入位置，不能替代分析仪自身身份。"
                ],
                [
                    ReportTable(
                        "每台设备证书准备摘要",
                        list(model.get("per_device_certificate_readiness") or []),
                    ),
                    ReportTable(
                        "每台设备点位证据",
                        _simple_rows(
                            model.get("per_device_point_evidence") or [],
                            (
                                "analyzer_device_id",
                                "component",
                                "calibratability_grade",
                                "fit_input_role",
                                "candidate_fit_allowed",
                                "candidate_fit_blockers",
                            ),
                            limit=240,
                        ),
                    ),
                ],
            ),
            ReportSection(
                "候选系数审核摘要",
                [
                    "候选系数只代表基于当前证据包的计算与审核入口，不代表已经写入设备。"
                    "每台分析仪按设备 ID 独立归并，blocker 必须处理后才能进入受控写入；warning 只进入审核提示。"
                ],
                [
                    ReportTable(
                        "候选系数按设备归并",
                        _simple_rows(
                            model.get("candidate_review_rollup") or [],
                            (
                                "component",
                                "analyzer_device_id",
                                "consolidated_status",
                                "blockers",
                                "warnings",
                                "review_action",
                            ),
                            limit=120,
                        ),
                    )
                ],
            ),
            ReportSection(
                "设备系数证据链",
                [
                    "GETCO1-9 快照、候选系数、写入事件和写后复验共同构成设备系数证据链。"
                    "任何 SENCOx 写入都必须能从旧系数、候选依据、写入命令、读回值和复验误差重建。"
                ],
                [
                    ReportTable(
                        "GETCO1-9 快照",
                        _simple_rows(
                            model.get("coefficient_snapshot_rows") or [],
                            (
                                "analyzer_id",
                                "snapshot_type",
                                "GETCO1",
                                "GETCO2",
                                "GETCO3",
                                "GETCO4",
                                "GETCO5",
                                "GETCO6",
                                "GETCO7",
                                "GETCO8",
                                "GETCO9",
                                "coefficients_hash",
                            ),
                            limit=120,
                        ),
                    ),
                    ReportTable(
                        "系数写入事件",
                        _simple_rows(
                            model.get("coefficient_write_event_rows") or [],
                            (
                                "analyzer_id",
                                "event_type",
                                "status",
                                "approved_by",
                                "command_summary",
                                "readback",
                            ),
                            limit=120,
                        ),
                    ),
                ],
            ),
            ReportSection(
                "CO2/H2O 结果",
                tables=[ReportTable("正式结果表", list(model.get("result_rows") or []))],
            ),
            ReportSection(
                "不确定度释放状态",
                [
                    f"不确定度状态：{(model.get('uncertainty_summary') or {}).get('status')}",
                    f"缺失必需输入：{';'.join((model.get('uncertainty_summary') or {}).get('missing_required') or [])}",
                ],
                [ReportTable("不确定度分量摘要", list((model.get("uncertainty_summary") or {}).get("component_summaries") or []))],
            ),
            ReportSection(
                "不确定度预算",
                [
                    "本表列出当前证据包可评估和未评估的不确定度分量。未释放的分量不得用于最终合格判定。",
                    "扩展不确定度 k=2 当前为 not_released，除非后续完成 GUM 预算和审核流程释放。",
                ],
                [ReportTable("不确定度分量", list(model.get("uncertainty_budget") or []))],
            ),
            ReportSection(
                "写后复验",
                [
                    "若本次存在 SENCO 写入，写后开放流通复验必须证明更新后的 CO2/H2O 输出仍与标准气或露点参考一致。",
                    f"写后复验状态：{post_write.get('overall_status')}",
                ],
                [
                    ReportTable(
                        "写后复验设备汇总",
                        _simple_rows(
                            post_write.get("device_summary") or [],
                            (
                                "device_id",
                                "component",
                                "point_count",
                                "pass_count",
                                "fail_count",
                                "max_abs_error",
                                "max_abs_error_pct",
                                "status",
                            ),
                        ),
                    ),
                    ReportTable(
                        "写后复验点位误差",
                        _simple_rows(
                            post_write.get("point_results") or [],
                            (
                                "device_id",
                                "component",
                                "point_id",
                                "standard_value",
                                "measured_value",
                                "error",
                                "error_pct",
                                "limit_value",
                                "status",
                                "reason",
                            ),
                            limit=200,
                        ),
                    ),
                ],
            ),
            ReportSection("系数写入声明", [model.get("coefficient_statement", "")]),
            ReportSection("限制和例外", list(model.get("limitations") or [])),
        ],
    )


def _device_slug(device_id: str) -> str:
    text = str(device_id or "unknown").strip() or "unknown"
    return "".join(ch if ch.isalnum() or ch in {"_", "-"} else "_" for ch in text)


def _rows_for_device(
    rows: Sequence[Mapping[str, Any]],
    device_id: str,
    *,
    keys: Sequence[str] = ("analyzer_device_id", "device_id", "analyzer_id", "sensor_id"),
) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for row in rows:
        if any(str(row.get(key) or "") == str(device_id) for key in keys):
            out.append(dict(row))
    return out


def _device_readiness_row(model: Mapping[str, Any], device_id: str) -> Dict[str, Any]:
    rows = _rows_for_device(
        model.get("per_device_certificate_readiness") or [],
        device_id,
        keys=("analyzer_device_id",),
    )
    return rows[0] if rows else {"analyzer_device_id": device_id}


def build_device_calibration_certificate(model: Mapping[str, Any], device_id: str) -> ReportDocument:
    release = model.get("report_release_decision") or {}
    readiness = _device_readiness_row(model, device_id)
    point_rows = _rows_for_device(
        model.get("per_device_point_evidence") or [],
        device_id,
        keys=("analyzer_device_id",),
    )
    candidate_rows = _rows_for_device(
        model.get("candidate_review_rollup") or [],
        device_id,
        keys=("analyzer_device_id",),
    )
    snapshot_rows = _rows_for_device(
        model.get("coefficient_snapshot_rows") or [],
        device_id,
        keys=("analyzer_id",),
    )
    write_rows = _rows_for_device(
        model.get("coefficient_write_event_rows") or [],
        device_id,
        keys=("analyzer_id",),
    )
    return ReportDocument(
        title=f"V1.5 气体分析仪校准证书（草稿） - {device_id}",
        sections=[
            ReportSection(
                "封面",
                [
                    f"报告编号：{model.get('report_no')}",
                    f"运行编号：{model.get('run_id')}",
                    f"被校分析仪设备 ID：{device_id}",
                    f"校准日期：{model.get('calibration_date')}",
                    f"校准地点：{model.get('location')}",
                    f"操作员：{model.get('operator')}",
                    f"审核员：{model.get('reviewer')}",
                    f"批准人：{model.get('approver')}",
                    f"签发状态：{readiness.get('calibration_certificate_status')}",
                    f"签发门禁：{release.get('release_status')} / {release.get('issue_mark')}",
                    f"门禁原因：{readiness.get('calibration_certificate_reasons')}",
                ],
            ),
            ReportSection(
                "校准范围和物理边界",
                [
                    model.get("scope_statement", ""),
                    "本证书草稿只汇总该设备 ID 的开放流通 CO2/H2O 主校准证据。"
                    "封路污染压力点、动态控压探针和 VENT-hold 不作为正式主拟合接受证据。",
                    model.get("pressure_compensation_statement", ""),
                ],
            ),
            ReportSection(
                "设备证书准备摘要",
                tables=[ReportTable("设备证书准备摘要", [readiness])],
            ),
            ReportSection(
                "点位可校准性",
                [
                    "A 级点可直接进入候选系数拟合；B 级点需要状态归一化或人工评审；"
                    "C 级点不得进入正式拟合。该表用于说明每个点为什么能用或不能用。"
                ],
                [
                    ReportTable(
                        "该设备点位证据",
                        _simple_rows(
                            point_rows,
                            (
                                "analyzer_device_id",
                                "component",
                                "calibratability_grade",
                                "fit_input_role",
                                "time_optimization_action",
                                "candidate_fit_allowed",
                                "candidate_fit_blockers",
                                "sample_readiness_status",
                                "sample_readiness_blockers",
                            ),
                            limit=240,
                        ),
                    )
                ],
            ),
            ReportSection(
                "候选系数和设备参数证据",
                [
                    "候选系数必须由该设备自身的 ratio/signal、压力、温度、露点和标准证书证据支持；"
                    "写入前后 GETCO/SENCO 状态必须可回溯。"
                ],
                [
                    ReportTable(
                        "候选系数评审摘要",
                        _simple_rows(
                            candidate_rows,
                            (
                                "component",
                                "analyzer_device_id",
                                "source_row_count",
                                "consolidated_status",
                                "blockers",
                                "warnings",
                                "review_action",
                            ),
                            limit=120,
                        ),
                    ),
                    ReportTable(
                        "GETCO1-9 快照",
                        _simple_rows(
                            snapshot_rows,
                            (
                                "analyzer_id",
                                "snapshot_type",
                                "GETCO1",
                                "GETCO2",
                                "GETCO3",
                                "GETCO4",
                                "GETCO5",
                                "GETCO6",
                                "GETCO7",
                                "GETCO8",
                                "GETCO9",
                                "coefficients_hash",
                            ),
                            limit=120,
                        ),
                    ),
                    ReportTable(
                        "系数写入事件",
                        _simple_rows(
                            write_rows,
                            (
                                "analyzer_id",
                                "event_type",
                                "status",
                                "candidate_id",
                                "approved_by",
                                "command_summary",
                                "readback",
                            ),
                            limit=120,
                        ),
                    ),
                ],
            ),
            ReportSection(
                "标准和参考设备",
                [
                    "标准气、压力、露点和温度参考证书共同支撑该设备本次测量结果的计量溯源。"
                ],
                [
                    ReportTable(
                        "标准气",
                        _simple_rows(
                            model.get("standard_gases") or [],
                            (
                                "component",
                                "cylinder_id",
                                "certificate_value",
                                "certificate_uncertainty",
                                "valid_until",
                                "supplier",
                                "certificate_hash",
                            ),
                        ),
                    ),
                    ReportTable(
                        "参考设备证书",
                        _simple_rows(
                            model.get("reference_certificates") or [],
                            ("reference_role", "certificate_id", "certificate_hash", "valid_until", "uncertainty", "unit"),
                        ),
                    ),
                ],
            ),
            ReportSection("限制和例外", list(model.get("limitations") or [])),
        ],
    )


def build_device_verification_certificate(model: Mapping[str, Any], device_id: str) -> ReportDocument:
    release = model.get("report_release_decision") or {}
    post_write = model.get("post_write_reverification") or {}
    readiness = _device_readiness_row(model, device_id)
    device_summary = _rows_for_device(
        post_write.get("device_summary") or [],
        device_id,
        keys=("device_id", "analyzer_device_id"),
    )
    point_results = _rows_for_device(
        post_write.get("point_results") or [],
        device_id,
        keys=("device_id", "analyzer_device_id"),
    )
    point_evidence = _rows_for_device(
        model.get("per_device_point_evidence") or [],
        device_id,
        keys=("analyzer_device_id",),
    )
    return ReportDocument(
        title=f"V1.5 气体分析仪检定/复验证书（草稿） - {device_id}",
        sections=[
            ReportSection(
                "封面",
                [
                    f"报告编号：{model.get('report_no')}",
                    f"运行编号：{model.get('run_id')}",
                    f"被校分析仪设备 ID：{device_id}",
                    f"复验状态：{readiness.get('verification_certificate_status')}",
                    f"复验门禁原因：{readiness.get('verification_certificate_reasons')}",
                    f"报告签发门禁：{release.get('release_status')} / {release.get('issue_mark')}",
                ],
            ),
            ReportSection(
                "检定/复验物理意义",
                [
                    "检定/复验不是重新拟合系数，而是在系数写入或候选评审后，"
                    "用独立开放流通点证明分析仪输出仍能回到标准气或湿度参考。"
                    "这一步用于防止离线模型正确但固件输出链路、SENCO 写入或水汽修正层异常。"
                ],
            ),
            ReportSection(
                "设备复验摘要",
                tables=[
                    ReportTable("证书准备摘要", [readiness]),
                    ReportTable("写后/独立复验设备汇总", list(device_summary)),
                ],
            ),
            ReportSection(
                "复验点误差",
                [
                    "该表必须保留每个复验点的标准值、测量值、误差、限值和判定。"
                    "若没有写后复验证据，本证书只能作为草稿或未覆盖声明。"
                ],
                [
                    ReportTable(
                        "复验点误差",
                        _simple_rows(
                            point_results,
                            (
                                "device_id",
                                "component",
                                "point_id",
                                "standard_value",
                                "measured_value",
                                "error",
                                "error_pct",
                                "limit_value",
                                "status",
                                "reason",
                            ),
                            limit=240,
                        ),
                    )
                ],
            ),
            ReportSection(
                "原始点位可校准性回看",
                [
                    "复验结论必须能追溯回主校准采样点；"
                    "若主校准点不是 A/B 级可解释点，复验通过也不能单独证明主校准链路合格。"
                ],
                [
                    ReportTable(
                        "主校准点位证据",
                        _simple_rows(
                            point_evidence,
                            (
                                "analyzer_device_id",
                                "component",
                                "calibratability_grade",
                                "fit_input_role",
                                "candidate_fit_allowed",
                                "candidate_fit_blockers",
                            ),
                            limit=240,
                        ),
                    )
                ],
            ),
        ],
    )


def _markdown_table(rows: Sequence[Mapping[str, Any]]) -> List[str]:
    if not rows:
        return ["_No rows._"]
    header: List[str] = []
    for row in rows:
        for key in row.keys():
            text = str(key)
            if text not in header:
                header.append(text)
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in rows:
        values = [str(row.get(key, "")).replace("|", "\\|").replace("\n", " ") for key in header]
        lines.append("| " + " | ".join(values) + " |")
    return lines


def render_markdown(document: ReportDocument) -> str:
    lines = [f"# {document.title}", ""]
    for section in document.sections:
        lines.extend([f"## {section.title}", ""])
        for paragraph in section.paragraphs:
            lines.extend([str(paragraph), ""])
        for table in section.tables:
            lines.extend([f"### {table.title}", ""])
            lines.extend(_markdown_table(table.rows))
            lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def _document_lines(document: ReportDocument) -> List[str]:
    lines = [document.title]
    for section in document.sections:
        lines.append(section.title)
        lines.extend(str(paragraph) for paragraph in section.paragraphs)
        for table in section.tables:
            lines.append(table.title)
            if table.rows:
                header = list(table.rows[0].keys())
                lines.append(" | ".join(header))
                for row in table.rows[:40]:
                    lines.append(" | ".join(str(row.get(key, "")) for key in header))
            else:
                lines.append("No rows.")
    return lines


def write_docx(document: ReportDocument, path: str | Path) -> Path:
    target = Path(path).resolve()
    target.parent.mkdir(parents=True, exist_ok=True)
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - python-docx is available in test env
        raise RuntimeError("python-docx is required for DOCX report output") from exc
    doc = Document()
    doc.add_heading(document.title, level=0)
    for section in document.sections:
        doc.add_heading(section.title, level=1)
        for paragraph in section.paragraphs:
            doc.add_paragraph(str(paragraph))
        for table in section.tables:
            doc.add_heading(table.title, level=2)
            if not table.rows:
                doc.add_paragraph("No rows.")
                continue
            header: List[str] = []
            for row in table.rows:
                for key in row.keys():
                    if key not in header:
                        header.append(str(key))
            grid = doc.add_table(rows=1, cols=len(header))
            grid.style = "Table Grid"
            for index, key in enumerate(header):
                grid.rows[0].cells[index].text = key
            for row in table.rows:
                cells = grid.add_row().cells
                for index, key in enumerate(header):
                    cells[index].text = str(row.get(key, ""))
    doc.save(target)
    return target


def _pdf_escape(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _write_minimal_pdf(lines: Sequence[str], path: Path) -> None:
    safe_lines = []
    for line in lines:
        encoded = str(line).encode("latin-1", errors="replace").decode("latin-1")
        safe_lines.append(encoded[:110])
    content = ["BT", "/F1 10 Tf", "50 780 Td"]
    first = True
    for line in safe_lines[:90]:
        if not first:
            content.append("0 -14 Td")
        content.append(f"({_pdf_escape(line)}) Tj")
        first = False
    content.append("ET")
    stream = "\n".join(content).encode("latin-1", errors="replace")
    objects: List[bytes] = []
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    objects.append(b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>")
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    objects.append(b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream")
    offsets: List[int] = []
    payload = bytearray(b"%PDF-1.4\n")
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(payload))
        payload.extend(f"{index} 0 obj\n".encode("ascii"))
        payload.extend(obj)
        payload.extend(b"\nendobj\n")
    xref = len(payload)
    payload.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode("ascii"))
    for offset in offsets:
        payload.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    payload.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(bytes(payload))


def write_pdf(document: ReportDocument, path: str | Path) -> Path:
    target = Path(path).resolve()
    target.parent.mkdir(parents=True, exist_ok=True)
    lines = _document_lines(document)
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfgen import canvas

        try:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            font_name = "STSong-Light"
        except Exception:
            font_name = "Helvetica"
        c = canvas.Canvas(str(target), pagesize=A4)
        width, height = A4
        y = height - 50
        c.setFont(font_name, 10)
        for line in lines:
            c.drawString(40, y, str(line)[:120])
            y -= 14
            if y < 40:
                c.showPage()
                c.setFont(font_name, 10)
                y = height - 50
        c.save()
    except Exception:
        _write_minimal_pdf(lines, target)
    return target


def write_report_document(document: ReportDocument, output_dir: str | Path, prefix: str) -> Dict[str, Path]:
    root = Path(output_dir).resolve()
    root.mkdir(parents=True, exist_ok=True)
    md_path = root / f"{prefix}.md"
    docx_path = root / f"{prefix}.docx"
    pdf_path = root / f"{prefix}.pdf"
    md_path.write_text(render_markdown(document), encoding="utf-8-sig")
    return {
        "markdown": md_path,
        "docx": write_docx(document, docx_path),
        "pdf": write_pdf(document, pdf_path),
    }


def write_v1_5_calibration_reports(
    *,
    evidence_bundle_path: str | Path,
    output_dir: str | Path,
    report_no: str = "",
    reviewer: str = "",
    approver: str = "",
    location: str = "",
    calibration_date: str = "",
    analyzer_prefix: str = "ga01",
    uncertainty_json: str | Path | None = None,
) -> Dict[str, Path]:
    bundle = _load_json(evidence_bundle_path)
    uncertainty_payload = _load_optional_json(uncertainty_json)
    model = build_report_model_from_bundle(
        bundle,
        report_no=report_no,
        reviewer=reviewer,
        approver=approver,
        location=location,
        calibration_date=calibration_date,
        analyzer_prefix=analyzer_prefix,
        uncertainty_payload=uncertainty_payload,
    )
    root = Path(output_dir).resolve()
    root.mkdir(parents=True, exist_ok=True)
    model_path = root / "report_model.json"
    model_path.write_text(json.dumps(model, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    outputs: Dict[str, Path] = {"report_model": model_path}
    for prefix, document in (
        ("run_report", build_run_report(model)),
        ("technical_report", build_technical_report(model)),
        ("formal_calibration_report", build_formal_calibration_report(model)),
    ):
        for key, path in write_report_document(document, root, prefix).items():
            outputs[f"{prefix}_{key}"] = path
    device_certificate_dir = root / "per_device_certificates"
    for row in model.get("per_device_certificate_readiness") or []:
        device_id = str(row.get("analyzer_device_id") or "").strip()
        if not device_id:
            continue
        slug = _device_slug(device_id)
        for prefix, document in (
            (f"device_{slug}_calibration_certificate", build_device_calibration_certificate(model, device_id)),
            (f"device_{slug}_verification_certificate", build_device_verification_certificate(model, device_id)),
        ):
            for key, path in write_report_document(document, device_certificate_dir, prefix).items():
                outputs[f"{prefix}_{key}"] = path
    outputs.update(_write_per_device_certificate_manifest(outputs=outputs, output_dir=root, model=model))
    return outputs
