from __future__ import annotations

import argparse
import math
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _fmt(value, digits: int = 3) -> str:
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except TypeError:
        pass
    if isinstance(value, float):
        if math.isinf(value):
            return ""
        if abs(value) >= 1000:
            return f"{value:.1f}"
        return f"{value:.{digits}f}"
    return str(value)


def _device_id(value) -> str:
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return ""
    try:
        return f"{int(value):03d}"
    except Exception:
        text = str(value)
        return text.zfill(3) if text.isdigit() else text


def _read_csv(base: Path, name: str) -> pd.DataFrame:
    path = base / name
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path)


def _records(df: pd.DataFrame, columns: Iterable[str], limit: int | None = None) -> list[dict[str, str]]:
    if df.empty:
        return []
    view = df.loc[:, [c for c in columns if c in df.columns]]
    if limit is not None:
        view = view.head(limit)
    rows: list[dict[str, str]] = []
    for row in view.to_dict("records"):
        rows.append({k: _fmt(v) for k, v in row.items()})
    return rows


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(8.5 if not bold else 9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]], title: str | None = None) -> None:
    if title:
        p = doc.add_paragraph()
        p.style = "Table Caption"
        p.add_run(title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        _set_cell_text(hdr_cells[i], header, bold=True)
        _shade_cell(hdr_cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            _set_cell_text(cells[i], value)
    doc.add_paragraph()


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _add_para(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Cm(0.74) if style is None else None
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(10.5)


def _set_document_style(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    for style_name in ["Normal", "Body Text"]:
        style = doc.styles[style_name]
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        style.font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption = doc.styles.add_style("Table Caption", 1)
    caption.font.name = "Microsoft YaHei"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.font.size = Pt(9)
    caption.font.bold = True


def _build_markdown(
    base: Path,
    health: pd.DataFrame,
    status: pd.DataFrame,
    residual: pd.DataFrame,
    snapshot: pd.DataFrame,
    flagged: pd.DataFrame,
    stale001: pd.DataFrame,
) -> str:
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    lines: list[str] = []
    lines.append("# 六台气体分析仪光学根因分析增强报告")
    lines.append("")
    lines.append(f"生成时间：{generated_at}")
    lines.append("")
    lines.append("数据来源：`v15_6ch_co2h2o_0612_r2` 历史 no-write 采样与离线分析证据。")
    lines.append("")
    lines.append("边界：本报告只读历史证据，不打开 COM，不写 SENCO，不控制气路、水路或压力控制器。")
    lines.append("")

    lines.append("## 一、结论摘要")
    lines.append("")
    lines.append("1. `079` 的主要风险是光学参考链路或归一化链路异常：ratio 短窗口可稳定，但 `ref_signal / CO2 signal / H2O signal` 与同点位正常组关系失真，误差随温度和点位翻转。")
    lines.append("2. `073` 不是普通拟合问题，而是 ratio/raw ratio 饱和或固件/量程/光学前端异常；不应进入组分系数写入。")
    lines.append("3. `001` 没有真正无效帧，主要是缓存偏旧；应通过程序的新鲜帧门禁和短恢复窗口规避。")
    lines.append("4. `091` 有低频缓存缺失/无效帧，优先按通信、缓存、主动上传节拍或串口稳定性处理，不应直接判为光路故障。")
    lines.append("5. 状态寄存器在原理上非常有价值，但本轮聚合证据里非空状态寄存器行数为 0；因此本轮不能用状态寄存器定位具体 bit，只能把它列为下一轮必须补齐的证据。")
    lines.append("")

    lines.append("## 二、六台设备健康总览")
    lines.append("")
    if not health.empty:
        cols = [
            "device_id", "point_count", "blocking_point_count", "max_relative_model_error_pct",
            "ref_signal_median", "ref_signal_range", "invalid_ratio_percent_max", "candidate_gate",
        ]
        lines.append("| 设备ID | 点数 | 阻断点 | 最大相对误差% | ref中位数 | ref范围 | 无效ratio% | 候选门禁 |")
        lines.append("| --- | ---: | ---: | ---: | ---: | ---: | ---: | --- |")
        for row in health[cols].to_dict("records"):
            lines.append(
                f"| {_device_id(row['device_id'])} | {_fmt(row['point_count'],0)} | {_fmt(row['blocking_point_count'],0)} | "
                f"{_fmt(row['max_relative_model_error_pct'],2)} | {_fmt(row['ref_signal_median'],1)} | "
                f"{_fmt(row['ref_signal_range'],1)} | {_fmt(row['invalid_ratio_percent_max'],2)} | {row['candidate_gate']} |"
            )
    lines.append("")

    lines.append("## 三、可能故障根因矩阵")
    cause_rows = [
        ["079", "高", "参考光路/参考探测通道温漂或污染", "同点位 ref_signal 偏离正常组，且残差符号随温度翻转", "检查参考光路、探测器、镜面污染；记录 ref_signal 随温度曲线"],
        ["079", "中-高", "SETCO2/SETILLUM 满值或 SETPOW 光源功率配置异常", "多点出现 ref_signal_near_configured_full_scale_hint", "读取 SETCO2/SETILLUM/SETPOW，与正常机和出厂记录比较"],
        ["079", "中", "固件归一化或信号链补偿异常", "ratio 稳定但 signal/ref 与误差不一致", "固件侧输出 signal/ref、raw ratio、报警 bit，并与软件离线模型比对"],
        ["079", "低", "标气或管路状态异常", "同一管路下其它设备无同类大误差", "仅作为反证保留，不作为主根因"],
        ["073", "高", "ratio/raw ratio 饱和或 ADC/量程溢出", "CO2 invalid ratio 95.45%，saturated_ratio_rows=416；ref_signal 低至约 45-70", "检查模拟前端、ADC量程、固件payload、光源/探测器连接"],
        ["073", "中-高", "参考信号链断路、严重污染或光源/接收器失效", "ref_signal_extremely_low_and_ratio_saturated", "暗场/亮场、SETPOW、参考通道输出专项测试"],
        ["091", "中", "主动上传帧缓存缺失或串口接收节拍问题", "CO2 缓存缺失 10 行，无效 ratio 2.27%，H2O 无不可用帧", "延长帧恢复窗口、记录 seq/age、检查串口线/供电/固件输出节拍"],
        ["001", "中", "采样锚点与 1Hz 主动上传帧时间不同步", "无不可用帧，但多点 stale_rows=10/10，age 约 0.9-1.0s", "采样前等待新帧；旧帧只允许进入诊断，不允许静默进入A级拟合"],
        ["077", "低-中", "显示层/S5-S6 或固件输出极值标记，不等同底层ratio故障", "底层 ratio 可用，但显示极值被标记", "复查输出层线性修正、显示上限/下限、S5/S6状态"],
        ["084", "低", "当前无明显光学链路异常", "工厂信号健康通过", "作为正常横向参考组保留"],
    ]
    lines.append("| 设备ID | 可能性 | 可能根因 | 支撑证据 | 建议验证 |")
    lines.append("| --- | --- | --- | --- | --- |")
    for row in cause_rows:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")

    lines.append("## 四、状态寄存器与无效帧分析")
    lines.append("")
    lines.append("状态寄存器有意义，但本轮证据不是“状态正常”，而是“状态寄存器未进入聚合证据”。它应在后续正式报告里作为独立证据字段：原始值、bit中文解释、对 FrameQC / FactorySignalQC / PointQC 的影响。")
    lines.append("")
    if not status.empty:
        lines.append("| 组分 | GA | 设备ID | 行数 | 非空状态行 | 不可用帧 | 缓存偏旧 | 无数据 | mode2 QC | 帧状态 |")
        lines.append("| --- | --- | --- | ---: | ---: | ---: | ---: | ---: | --- | --- |")
        for row in status.to_dict("records"):
            lines.append(
                f"| {row['component']} | {row['ga']} | {_device_id(row['device_id'])} | {_fmt(row['rows'],0)} | "
                f"{_fmt(row['status_nonblank_rows'],0)} | {_fmt(row['frame_unusable_rows'],0)} | "
                f"{_fmt(row['frame_stale_rows'],0)} | {_fmt(row['frame_has_no_data_rows'],0)} | "
                f"{row['mode2_qc_counts']} | {row['frame_status_counts']} |"
            )
    lines.append("")

    lines.append("### 4.1 001 的程序规避")
    lines.append("")
    lines.append("001 的问题是 stale frame，不是解析失败。程序上应执行：采样窗口取帧前先等待当前设备新帧；如果帧龄仍超过阈值，则标记为 stale，降级或拒绝该设备该窗口；不能把上一秒缓存帧直接当作当前开路流通状态。")
    if not stale001.empty:
        lines.append("")
        lines.append("| 组分 | 点位 | 温度 | 目标 | 行数 | stale行 | 最大帧龄ms |")
        lines.append("| --- | --- | ---: | ---: | ---: | ---: | ---: |")
        for row in stale001.head(8).to_dict("records"):
            lines.append(
                f"| {row['component']} | {row['point_tag']} | {_fmt(row['temp_set_c'],0)} | {_fmt(row['target_value'],2)} | "
                f"{_fmt(row['rows'],0)} | {_fmt(row['stale_rows'],0)} | {_fmt(row['max_age_ms'],1)} |"
            )
    lines.append("")

    lines.append("### 4.2 091 的程序规避")
    lines.append("")
    lines.append("091 的 CO2 有 10 行缓存缺失。程序上应按单设备恢复处理：先等待一个短恢复窗口；仍无新帧时，只把 091 的该点标为 missing/stale，不拖死其它设备，也不把缺帧补成平均值。复验和报告中应显示缺帧比例、发生点位和是否影响拟合等级。")
    lines.append("")

    lines.append("## 五、079 重点根因证据")
    if not flagged.empty:
        flagged079 = flagged[flagged["device_id"].astype(str).isin(["79", "079"])].copy()
        flagged079["abs_rel"] = flagged079["relative_error_pct"].abs()
        flagged079 = flagged079.sort_values(["blocks_candidate_write", "abs_rel"], ascending=[False, False]).head(12)
        lines.append("")
        lines.append("| 组分 | 点位 | T | 目标 | ratio跨度 | ref | CO2信号 | H2O信号 | 误差 | 相对误差% | 标志 |")
        lines.append("| --- | --- | ---: | ---: | ---: | ---: | ---: | ---: | ---: | ---: | --- |")
        for row in flagged079.to_dict("records"):
            lines.append(
                f"| {row['component']} | {row['point_tag']} | {_fmt(row['temp_set_c'],0)} | {_fmt(row['target'],2)} | "
                f"{_fmt(row['ratio_span'],4)} | {_fmt(row['ref_signal_median'],1)} | {_fmt(row['co2_signal_median'],1)} | "
                f"{_fmt(row['h2o_signal_median'],1)} | {_fmt(row['model_error'],2)} | {_fmt(row['relative_error_pct'],2)} | "
                f"{row['signal_health_flags']} |"
            )
    lines.append("")
    lines.append("判断：079 不是因为某一瓶气或某一个温度点异常。它在多个温度、多个浓度点出现 `stable_ratio_but_reference_chain_unhealthy`，说明底层光学归一化链路比拟合算法更可疑。")
    lines.append("")

    lines.append("## 六、073 重点根因证据")
    if not snapshot.empty:
        snap073 = snapshot[snapshot["device_id"].astype(str).isin(["73", "073"])].copy()
        lines.append("")
        lines.append("| 组分 | 行数 | 无效ratio | 饱和ratio | 无效ratio% | ratio中位 | raw ratio最大 | ref中位 | ref最小 | 物理解释 |")
        lines.append("| --- | ---: | ---: | ---: | ---: | ---: | ---: | ---: | ---: | --- |")
        for row in snap073.to_dict("records"):
            lines.append(
                f"| {row['component']} | {_fmt(row['rows'],0)} | {_fmt(row['invalid_ratio_rows'],0)} | "
                f"{_fmt(row['saturated_ratio_rows'],0)} | {_fmt(row['invalid_ratio_percent'],2)} | {_fmt(row['ratio_median'],4)} | "
                f"{_fmt(row['raw_ratio_max'],2)} | {_fmt(row['ref_signal_median'],1)} | {_fmt(row['ref_signal_min'],1)} | "
                f"{row['physical_interpretation']} |"
            )
    lines.append("")

    lines.append("## 七、对仪器硬件/固件的建议")
    lines.append("")
    lines.append("1. 固件应输出并报警：ref_signal 过低/过高、signal/ref 异常、ratio/raw ratio 饱和、脉冲不同步、光功率异常、光电流异常。")
    lines.append("2. 对 079 优先检查 SETCO2/SETILLUM/SETPOW、参考光路、探测器增益、温漂和固件归一化，不建议继续用组分系数吸收错误。")
    lines.append("3. 对 073 优先检查光源/接收器/ADC/固件payload，不建议进入正式拟合。")
    lines.append("4. 对 091 优先改善串口帧完整性、主动上传节拍、缓存刷新与程序恢复窗口。")
    lines.append("5. 对 001 优先优化采样锚点和帧新鲜度门禁，不按硬件故障处理。")
    lines.append("")

    lines.append("## 八、对 V1.5 程序的建议")
    lines.append("")
    lines.append("1. 采样前必须确认阀门仍处于开路流通状态，且每台设备拿到当前窗口的新鲜帧。")
    lines.append("2. 建立 `FactorySignalHealthGate`：只要 ref_signal 或 signal/ref 关系阻断，就不允许写 S1/S2/S3/S4。")
    lines.append("3. 状态寄存器应作为独立证据：缺失、正常、异常三态都要写入报告。")
    lines.append("4. 某台设备无效帧、缓存缺失或 ratio 异常，只影响该设备该点，不应污染其它设备。")
    lines.append("5. 报告中必须区分：光学链路故障、串口/缓存帧故障、输出层显示极值、气体/管路未稳定。")
    lines.append("")

    lines.append("## 九、附件清单")
    for name in [
        "six_device_optical_health_summary_extended.csv",
        "same_point_peer_comparison.csv",
        "factory_signal_flagged_points.csv",
        "component_signal_snapshot.csv",
        "status_register_and_invalid_frame_summary.csv",
        "invalid_frame_examples.csv",
        "id001_stale_frame_point_breakdown.csv",
    ]:
        lines.append(f"- `{base / name}`")
    lines.append("")
    return "\n".join(lines)


def _add_markdown_like_content(doc: Document, markdown: str) -> None:
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            _add_heading(doc, line[2:], 0)
        elif line.startswith("## "):
            _add_heading(doc, line[3:], 1)
        elif line.startswith("### "):
            _add_heading(doc, line[4:], 2)
        elif line.startswith("- "):
            _add_para(doc, "• " + line[2:])
        elif line[:3].strip().endswith(".") and line[:1].isdigit():
            _add_para(doc, line)
        elif line.startswith("|"):
            continue
        else:
            _add_para(doc, line)


def _docx_report(
    out_docx: Path,
    markdown: str,
    health: pd.DataFrame,
    status: pd.DataFrame,
    snapshot: pd.DataFrame,
    flagged: pd.DataFrame,
    stale001: pd.DataFrame,
) -> None:
    doc = Document()
    _set_document_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("六台气体分析仪光学根因分析增强报告")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("V1.5 no-write 采样证据 / 工厂模式信号 / 帧质量 / 状态寄存器证据")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)

    _add_para(doc, f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    _add_para(doc, "边界说明：本报告只读历史证据，没有打开 COM，没有写 SENCO，没有控制气路、水路或压力控制器。")
    doc.add_page_break()

    _add_heading(doc, "一、核心结论", 1)
    conclusions = [
        "079 是当前最值得追查的光学异常设备：ratio 短窗口稳定，但 ref_signal / CO2 signal / H2O signal 与同点位正常组关系失真，说明参考链路或归一化链路风险高。",
        "073 呈现 ratio/raw ratio 饱和和参考信号极低特征，优先怀疑量程、ADC、光学前端、参考链路或固件 payload 合同异常，不应进入普通组分系数拟合。",
        "001 的问题不是坏帧，而是缓存偏旧帧；程序应等待当前窗口新鲜帧，旧帧不能静默进入 A 级拟合。",
        "091 的问题是低频缓存缺失/无效帧，优先从串口接收、主动上传节拍、缓存刷新和固件帧完整性解决。",
        "状态寄存器具有重要诊断意义，但本轮聚合证据中状态寄存器为空；下一轮应强制归档原始状态字和中文 bit 解释。",
    ]
    for item in conclusions:
        _add_para(doc, item)

    if not health.empty:
        _add_heading(doc, "二、六台设备健康总览", 1)
        rows = []
        for row in health.to_dict("records"):
            rows.append([
                _device_id(row.get("device_id")),
                _fmt(row.get("point_count"), 0),
                _fmt(row.get("blocking_point_count"), 0),
                _fmt(row.get("max_relative_model_error_pct"), 2),
                _fmt(row.get("ref_signal_median"), 1),
                _fmt(row.get("ref_signal_range"), 1),
                _fmt(row.get("invalid_ratio_percent_max"), 2),
                str(row.get("candidate_gate", "")),
            ])
        _add_table(doc, ["设备ID", "点数", "阻断点", "最大相对误差%", "ref中位数", "ref范围", "无效ratio%", "候选门禁"], rows)

    _add_heading(doc, "三、可能故障根因矩阵", 1)
    root_rows = [
        ["079", "高", "参考光路/参考探测通道温漂或污染", "同点位 ref_signal 偏离正常组，残差随温度翻转", "检查参考光路、探测器、镜面污染；记录 ref_signal 随温度曲线"],
        ["079", "中-高", "SETCO2/SETILLUM 满值或 SETPOW 光源功率配置异常", "多点出现 ref_signal_near_configured_full_scale_hint", "读取 SETCO2/SETILLUM/SETPOW，与正常机比较"],
        ["073", "高", "ratio/raw ratio 饱和或 ADC/量程溢出", "CO2 invalid ratio 95.45%，ref_signal 极低", "检查 ADC、光源、接收器、payload合同"],
        ["091", "中", "主动上传帧缓存缺失或串口接收节拍问题", "CO2 缓存缺失 10 行，H2O 无不可用帧", "增加新鲜帧恢复窗口，检查串口线/供电/固件节拍"],
        ["001", "中", "采样锚点与 1Hz 主动上传帧不同步", "多点 stale_rows=10/10，最大帧龄约 1s", "采样前等待新帧；旧帧降级或拒绝"],
        ["077", "低-中", "显示层/S5-S6 或固件输出极值标记", "底层 ratio 可用但显示极值标记", "复查输出层线性修正和显示上下限"],
        ["084", "低", "当前无明显光学链路异常", "工厂信号健康通过", "作为正常横向参考保留"],
    ]
    _add_table(doc, ["设备ID", "可能性", "可能根因", "支撑证据", "建议验证"], root_rows)

    _add_heading(doc, "四、状态寄存器与无效帧分析", 1)
    _add_para(doc, "状态寄存器能帮助区分光功率异常、光电流异常、脉冲不同步、CO2/H2O 信号超限、温度异常和数据变化量超标等内部故障。本轮 CSV 中 status_nonblank_rows 全为 0，因此不能把它解释为状态正常，只能判定为状态寄存器证据缺失。")
    if not status.empty:
        rows = []
        for row in status.to_dict("records"):
            rows.append([
                str(row.get("component", "")),
                str(row.get("ga", "")),
                _device_id(row.get("device_id")),
                _fmt(row.get("rows"), 0),
                _fmt(row.get("status_nonblank_rows"), 0),
                _fmt(row.get("frame_unusable_rows"), 0),
                _fmt(row.get("frame_stale_rows"), 0),
                _fmt(row.get("frame_has_no_data_rows"), 0),
                str(row.get("frame_status_counts", "")),
            ])
        _add_table(doc, ["组分", "GA", "设备ID", "行数", "非空状态", "不可用帧", "缓存偏旧", "无数据", "帧状态"], rows)

    _add_heading(doc, "五、001 与 091 的程序规避方案", 1)
    _add_para(doc, "001：把缓存偏旧帧作为时间对齐问题处理。采样窗口开始前等待当前设备的新鲜帧；若仍旧帧，标为 stale，不能进入 A 级拟合。该策略避免把上一个气体物理状态混入当前标准气窗口。")
    _add_para(doc, "091：把缓存缺失作为单台设备帧质量问题处理。程序应给 091 一个短恢复窗口；恢复失败时只降级或拒绝 091 的该点，不拖死其它设备，不用插值替代真实帧。")
    if not stale001.empty:
        rows = []
        for row in stale001.head(8).to_dict("records"):
            rows.append([
                str(row.get("component", "")),
                str(row.get("point_tag", "")),
                _fmt(row.get("temp_set_c"), 0),
                _fmt(row.get("target_value"), 2),
                _fmt(row.get("stale_rows"), 0),
                _fmt(row.get("max_age_ms"), 1),
            ])
        _add_table(doc, ["组分", "点位", "温度", "目标", "stale行", "最大帧龄ms"], rows, "001 缓存偏旧帧代表点位")

    _add_heading(doc, "六、079 重点证据", 1)
    _add_para(doc, "079 在多个温度与多个标气点出现“ratio 稳定但参考链路不健康”。这说明问题不是单一气瓶、单一温度点或单一压力输入，而是参考链路、光源/探测器、满值配置或固件归一化的组合风险。")
    if not flagged.empty:
        f079 = flagged[flagged["device_id"].astype(str).isin(["79", "079"])].copy()
        f079["abs_rel"] = f079["relative_error_pct"].abs()
        f079 = f079.sort_values(["blocks_candidate_write", "abs_rel"], ascending=[False, False]).head(10)
        rows = []
        for row in f079.to_dict("records"):
            rows.append([
                str(row.get("component", "")),
                str(row.get("point_tag", "")),
                _fmt(row.get("temp_set_c"), 0),
                _fmt(row.get("target"), 2),
                _fmt(row.get("ratio_span"), 4),
                _fmt(row.get("ref_signal_median"), 1),
                _fmt(row.get("model_error"), 2),
                _fmt(row.get("relative_error_pct"), 2),
                str(row.get("signal_health_flags", "")),
            ])
        _add_table(doc, ["组分", "点位", "T", "目标", "ratio跨度", "ref", "误差", "相对误差%", "标志"], rows)

    _add_heading(doc, "七、073 重点证据", 1)
    if not snapshot.empty:
        s073 = snapshot[snapshot["device_id"].astype(str).isin(["73", "073"])].copy()
        rows = []
        for row in s073.to_dict("records"):
            rows.append([
                str(row.get("component", "")),
                _fmt(row.get("rows"), 0),
                _fmt(row.get("invalid_ratio_rows"), 0),
                _fmt(row.get("saturated_ratio_rows"), 0),
                _fmt(row.get("invalid_ratio_percent"), 2),
                _fmt(row.get("ratio_median"), 4),
                _fmt(row.get("raw_ratio_max"), 2),
                _fmt(row.get("ref_signal_median"), 1),
                _fmt(row.get("ref_signal_min"), 1),
            ])
        _add_table(doc, ["组分", "行数", "无效ratio", "饱和ratio", "无效ratio%", "ratio中位", "raw最大", "ref中位", "ref最小"], rows)
    _add_para(doc, "073 的 ref_signal 极低且 ratio 饱和，不应通过改 S1/S2/S3/S4 尝试修正。应先检查硬件光学链路、ADC量程、固件payload和满值配置。")

    _add_heading(doc, "八、维修与软件改进清单", 1)
    actions = [
        "079：读取 SETCO2/SETILLUM/SETPOW，检查参考光路、光源、探测器、模拟前端和温漂；完成前不要写组分主系数。",
        "073：检查固件版本、payload合同、ADC/量程、参考通道和光源/探测器；不进入普通校准拟合。",
        "091：优化主动上传接收缓存、帧新鲜度恢复窗口、串口间隔和缺帧统计。",
        "001：把旧帧排除出 A 级样本；采样前等待新帧并记录 frame_age_ms。",
        "V1.5：正式报告增加状态寄存器原始值、中文 bit 解释、FactorySignalHealthGate 和每台设备独立评级。",
    ]
    for action in actions:
        _add_para(doc, action)

    _add_heading(doc, "九、附件", 1)
    for name in [
        "same_point_peer_comparison.csv",
        "factory_signal_flagged_points.csv",
        "component_signal_snapshot.csv",
        "status_register_and_invalid_frame_summary.csv",
        "invalid_frame_examples.csv",
        "id001_stale_frame_point_breakdown.csv",
    ]:
        _add_para(doc, str(out_docx.parent / name))

    doc.save(out_docx)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--evidence-dir", required=True, type=Path)
    args = parser.parse_args()
    base = args.evidence_dir

    health = _read_csv(base, "six_device_optical_health_summary_extended.csv")
    status = _read_csv(base, "status_register_and_invalid_frame_summary.csv")
    residual = _read_csv(base, "device_component_residual_signal_summary.csv")
    snapshot = _read_csv(base, "component_signal_snapshot.csv")
    flagged = _read_csv(base, "factory_signal_flagged_points.csv")
    stale001 = _read_csv(base, "id001_stale_frame_point_breakdown.csv")

    markdown = _build_markdown(base, health, status, residual, snapshot, flagged, stale001)
    out_md = base / "six_device_optical_root_cause_report_zh_enhanced.md"
    out_docx = base / "六台气体分析仪光学根因分析增强报告.docx"
    out_md.write_text(markdown, encoding="utf-8")
    _docx_report(out_docx, markdown, health, status, snapshot, flagged, stale001)

    with zipfile.ZipFile(out_docx, "r") as zf:
        required = {"[Content_Types].xml", "word/document.xml"}
        missing = required - set(zf.namelist())
        if missing:
            raise RuntimeError(f"DOCX package missing required parts: {sorted(missing)}")

    print(f"markdown={out_md}")
    print(f"docx={out_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
